'''
print_patient.py
----------------
Service for printing all patient visits information.

inputs: patient (dict)
outputs: Word document file stream (.docx) and filename (str)
'''

import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import session

def create_patient_document(patient):
    # Create a new Document
    doc = Document()

    # Add the doctor's name as the title, bold and centered
    doctor_name = patient.get("dr", "Unknown Doctor")
    title = doc.add_heading(level=1)
    run = title.add_run(f"{doctor_name}")
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    patient_no = session['patientno']+1

    # Add patient details
    doc.add_heading(f"Patient Name: {patient.get('name', 'Unknown')}", level=2)
    doc.add_paragraph(f"ID No.: {patient.get('id', ' ')}")
    doc.add_paragraph(f"Patient No.: {patient_no}")
    doc.add_paragraph(f"Location: {patient.get('location', ' ')}")
    doc.add_paragraph(f"Age: {patient.get('age', ' ')}")
    doc.add_paragraph(f"Phone: {patient.get('phone', ' ')}")
    doc.add_paragraph(f"Past Medical History: {patient.get('pmh', ' ')}")
    doc.add_paragraph(f"Allergies: {patient.get('allergies', ' ')}")
    nnn=patient.get('next', ' ')
    if nnn == "<built-in method date of datetime.datetime object at 0x79d2fbd5a700>":
        nnn=' '
    doc.add_paragraph(f"Next visit: {nnn}")

    for visit in patient.get("visits", []):
        table = doc.add_table(rows=6, cols=2)

        table.style = 'Table Grid'
        cells = table.rows[0].cells

        cells[0].text = 'Visit No:'
        cells[1].text = str(visit.get('vno', 'N/A'))
        cells = table.rows[1].cells

        cells[0].text = 'Visit Date:'
        cells[1].text = visit.get('visit_date', 'N/A')
        cells = table.rows[2].cells

        cells[0].text = 'Diagnosis:'
        cells[1].text = visit.get('diagnosis', 'N/A')
        cells = table.rows[3].cells

        cells[0].text = 'Notes:'
        cells[1].text = str(visit.get('details', 'N/A'))
        cells = table.rows[4].cells

        cells[0].text = 'Lab Results:'
        cells[1].text = visit.get('lab', 'N/A')
        cells = table.rows[5].cells

        cells[0].text = 'Treatment:'
        cells[1].text = visit.get('treatment', 'N/A')

        # Add a line break after each visit
        doc.add_paragraph("\n")

    # Save the document to an in-memory file-like object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    n=patient.get('name', 'Unknown')
    return file_stream, f"{n.replace(' ', '_')}_data.docx"