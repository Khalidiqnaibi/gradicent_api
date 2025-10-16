import os
import io  # Import the io module
import shutil
import pathlib
import json
import html
import requests
from cryptography.fernet import Fernet, InvalidToken
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
import base64
from flask import sessions, Flask, jsonify, render_template, redirect, request, session, url_for, abort, send_file,send_from_directory 
from google.oauth2 import id_token
from datetime import datetime,timedelta
from google_auth_oauthlib.flow import Flow
from pip._vendor import cachecontrol
import google.auth.transport.requests
import webbrowser
import firebase_admin
from firebase_admin import credentials, db, firestore, storage
from google.cloud.firestore_v1.base_query import FieldFilter
from werkzeug.utils import secure_filename
from io import BytesIO
import paypalrestsdk
from paypalrestsdk import Payment
import secrets
import uuid
import string
from zenora import APIClient
import logging
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from time import sleep
from urllib.parse import quote
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import uuid

##############
#~ Immortal ~#
##############

appname='Binder'
cname='RiaSoftware'
app = Flask(f"{cname} Api",static_url_path='/home/RiaSoftware/s/static',template_folder='/home/RiaSoftware/s/templates')


app.secret_key = "abcdefghijk123"
appurl='http://www.bindersoftware.com'

basicprice=22
premprice=300
medprice=215

# Binder Medical
starter_price = 5
pro_price = 25
ultra_price = 125

global bPay
bPay =False
global closee
closee =False
global prePay
prePay= False
global password
password = "@Ksoftkhaafif1"  # Change this to your secure password
ADMIN_SECRET = "bindersoftware.com"  # change to env var in production

PAYPAL_CLIENT_ID ="AaH6jy2wDk69MEKa5aVYIwz06AMJwjym3qziA3wmF0qlbdKtcI-iIZCmj9qjK2mcHvEXgXbVnyq_6nP1",
PAYPAL_SECRET = "EFFgbsOtPSMXBbRyivM5ogXekW4BMETUkjcBJf9LCMRuWGaqxTtAVOWHa30WkwP-w19eQ6b8aMHIxFf9"

# Initialize Firebase with your Firebase project's credentials
cred = credentials.Certificate(r"/home/RiaSoftware/s/key2.json")
firebase_admin.initialize_app(cred, {
    'databaseURL': 'https://monydb-f2cdb-default-rtdb.europe-west1.firebasedatabase.app/',
    'storageBucket': 'monydb-f2cdb.appspot.com'
})
os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

PADDLE_API_KEY = "apikey_01k4wa7wtjxbnv8jwy6hzw6g5e"

GOOGLE_CLIENT_ID = "107932074863-nlil9n5j9lmahqfb15cmn52u59evpse9.apps.googleusercontent.com"
client_secrets_file = os.path.join(pathlib.Path(__file__).parent, r"/home/RiaSoftware/s/client_secret1.json")

flow = Flow.from_client_secrets_file(
    client_secrets_file=client_secrets_file,
    scopes=["https://www.googleapis.com/auth/userinfo.profile", "https://www.googleapis.com/auth/userinfo.email", "openid"],
    redirect_uri=f"{appurl}/callback"
)
def nameee():
    session['appname']=appname

def create_fernet():
    global password
    password = password.encode()  # Convert password to bytes
    salt = b"MY_Potatolikesme_too"  # Change this to a unique salt
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        iterations=100000,  # Adjust the number of iterations as needed
        salt=salt,
        length=32  # The length of the derived key
    )
    key = base64.urlsafe_b64encode(kdf.derive(password))  # Derive key
    return Fernet(key)

def encrypt_data(data):
    if isinstance(data, int):  # Convert integers to string before encryption
        data = str(data)
    return fernet.encrypt(data.encode())  # Convert to bytes and encrypt

def decrypt_data(encrypted_data):
    try:
        if isinstance(encrypted_data, str):
            encrypted_data = encrypted_data.encode()  # Convert to bytes if necessary
        decrypted_data = fernet.decrypt(encrypted_data).decode()  # Decrypt and decode to string

        # Try to convert back to integer if it was originally an integer
        if decrypted_data.isdigit():
            return int(decrypted_data)

        return decrypted_data  # Return the string if it's not an integer
    except InvalidToken:
        # Handle decryption failure (e.g., incorrect password or corrupted data)
        return None
    except TypeError as e:
        # Handle errors due to incorrect types
        print(f"TypeError: {e}")
        return None

def load_user_data():
        path = f"C:/{cname}/{appname}/user_data.txt"

        user_data = {
            "first": '2001-01-01',
            "google_id": 1,
            "name": "non",
            "payed": 0
        }

        if not os.path.exists(f"C:/{cname}"):
            os.makedirs(f"C:/{cname}")

        if not os.path.exists(f"C:/{cname}/{appname}"):
            os.makedirs(f"C:/{cname}/{appname}")

        if path and os.path.exists(path):
            with open(path, "rb") as json_file:
                encrypted_data = json_file.read()
                user_data = decrypt_data(encrypted_data)

        return user_data

def load_app_data():
        app_data_path = f"C:/{cname}/{appname}/data.txt"

        if not os.path.exists(f"C:/{cname}"):
            os.makedirs(f"C:/{cname}")

        if not os.path.exists(f"C:/{cname}/{appname}"):
            os.makedirs(f"C:/{cname}/{appname}")

        if os.path.exists(app_data_path):
            with open(app_data_path, 'rb') as file:
                encrypted_data = file.read()
                app_data = decrypt_data(encrypted_data)
        else:
            save_app_data()

        apiurl = app_data.get('url', '')  # Assuming 'url' might not always be present

        return app_data

def login_is_required(function):
    def wrapper(*args, **kwargs):
        if "google_id" not in session:
            return abort(401)
        return function(*args, **kwargs)
    wrapper.__name__ = function.__name__  # force Flask to see the right name
    return wrapper

def admin_required(func):
    def wrapper(*args, **kwargs):
        key = request.args.get("key") or request.headers.get("X-Admin-Key")
        if key != ADMIN_SECRET:
            return jsonify({"error": "Unauthorized"}), 403
        return func(*args, **kwargs)
    wrapper.__name__ = func.__name__
    return wrapper

def iscode(code):
    # Reference to the Firebase Realtime Database
    ref = db.reference('/codes')
    # Query the database to find the user with the matching google_id
    query = ref.order_by_child('code').equal_to(code).limit_to_first(1)

    # Execute the query
    result = query.get()

    if not result:
        # User with the provided google_id not found
        return None

    user_id, user_data = next(iter(result.items()))

    session['cd']=user_data

    if user_id and not user_data['used']:
        return user_data['plan']
    else:
        return None

def quittab():
    # Get a list of open browser tabs
    browser_tabs = webbrowser.get().tab_list()

    # Check if any tabs match the condition
    for tab in browser_tabs:
        if tab.url.endswith('/protected_area'):
            # Close the tab if the URL ends with '/protected_area'
            tab.close()

def add_user(user_data):
    # Reference to the Firebase Realtime Database
    ref = db.reference('/users')  # '/users' is the path where you want to store user data

    # Push user data to the database
    new_user_ref = ref.push(user_data)

def save_code(code,plan,gid=""):
    # Reference to the Firebase Realtime Database
    ref = db.reference('/codes')  # '/codes' is the path where you want to store codes ref.order_by_child('code').equal_to(code)

    v= {"code":code,"date":datetime.now().date().isoformat(),'plan':plan,"used":False,"google_id":gid}
    # Push the code to the database

    ref.child(code).set(v)
    #new_code_ref = ref.push(v)

def save_seccode(code,gid=""):

    v= {"code":code,"date":datetime.now().date().isoformat(),'plan':'sec',"users":0,"google_id":gid}
    typee=session.get('wtype', 'drs')
    dr_ref = db.reference(f'/{typee}/{gid}')
    nn=dr_ref.get()
    if 'settings' not in nn:
        nn["settings"]={'ac':v}
    else:
        nn["settings"]['ac']=v
    dr_ref.update(nn)

def generate_strong_password(length=12):
    characters = string.ascii_letters + string.digits
    password = ''.join(secrets.choice(characters) for _ in range(length))
    return password

def gencode():
    code=f"{generate_strong_password(4)}-{generate_strong_password(4)}-{generate_strong_password(4)}"
    return code

def create_patient_document(patient):
    # Create a new Document
    doc = Document()

    # Add the doctor's name as the title, bold and centered
    doctor_name = patient.get("dr", "Unknown Doctor")
    title = doc.add_heading(level=1)
    run = title.add_run(f"{doctor_name}")
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    patient_no = session['patientno']+1

    # Add patient details
    doc.add_heading(f"Patient Name: {patient.get('name', 'Unknown')}", level=2)
    doc.add_paragraph(f"ID No.: {patient.get('id', ' ')}")
    doc.add_paragraph(f"Patient No.: {patient_no}")
    doc.add_paragraph(f"Location: {patient.get('location', ' ')}")
    doc.add_paragraph(f"Age: {patient.get('age', ' ')}")
    doc.add_paragraph(f"Phone: {patient.get('phone', ' ')}")
    doc.add_paragraph(f"Past Medical History: {patient.get('pmh', ' ')}")
    doc.add_paragraph(f"Allergies: {patient.get('allergies', ' ')}")
    nnn=patient.get('next', ' ')
    if nnn == "<built-in method date of datetime.datetime object at 0x79d2fbd5a700>":
        nnn=' '
    doc.add_paragraph(f"Next visit: {nnn}")

    for visit in patient.get("visits", []):
        table = doc.add_table(rows=6, cols=2)

        table.style = 'Table Grid'
        cells = table.rows[0].cells

        cells[0].text = 'Visit No:'
        cells[1].text = str(visit.get('vno', 'N/A'))
        cells = table.rows[1].cells

        cells[0].text = 'Visit Date:'
        cells[1].text = visit.get('visit_date', 'N/A')
        cells = table.rows[2].cells

        cells[0].text = 'Diagnosis:'
        cells[1].text = visit.get('diagnosis', 'N/A')
        cells = table.rows[3].cells

        cells[0].text = 'Notes:'
        cells[1].text = str(visit.get('details', 'N/A'))
        cells = table.rows[4].cells

        cells[0].text = 'Lab Results:'
        cells[1].text = visit.get('lab', 'N/A')
        cells = table.rows[5].cells

        cells[0].text = 'Treatment:'
        cells[1].text = visit.get('treatment', 'N/A')

        # Add a line break after each visit
        doc.add_paragraph("\n")

    # Save the document to an in-memory file-like object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    n=patient.get('name', 'Unknown')
    return file_stream, f"{n.replace(' ', '_')}_data.docx"

def create_visit_document(patient,vno):
    # Create a new Document
    doc = Document()

    # Add the doctor's name as the title, bold and centered
    doctor_name = patient.get("dr", "Unknown Doctor")
    title = doc.add_heading(level=1)
    run = title.add_run(f"{doctor_name}")
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    patient_no = session['patientno']+1

    # Add patient details
    doc.add_heading(f"Patient Name: {patient.get('name', 'Unknown')}", level=2)
    doc.add_paragraph(f"ID No.: {patient.get('id', ' ')}")
    doc.add_paragraph(f"Patient No.: {patient_no}")
    doc.add_paragraph(f"Location: {patient.get('location', ' ')}")
    doc.add_paragraph(f"Age: {patient.get('age', ' ')}")
    doc.add_paragraph(f"Phone: {patient.get('phone', ' ')}")
    doc.add_paragraph(f"Past Medical History: {patient.get('pmh', ' ')}")
    doc.add_paragraph(f"Allergies: {patient.get('allergies', ' ')}")

    nnn=patient.get('next', ' ')
    if nnn == "<built-in method date of datetime.datetime object at 0x79d2fbd5a700>":
        nnn=' '
    doc.add_paragraph(f"Next visit: {nnn}")

    visit =patient["visits"][vno]

    if visit :
        table = doc.add_table(rows=6, cols=2)

        table.style = 'Table Grid'
        cells = table.rows[0].cells

        cells[0].text = 'Visit No:'
        cells[1].text = str(visit.get('vno', 'N/A'))
        cells = table.rows[1].cells

        cells[0].text = 'Visit Date:'
        cells[1].text = visit.get('visit_date', 'N/A')
        cells = table.rows[2].cells

        cells[0].text = 'Diagnosis:'
        cells[1].text = visit.get('diagnosis', 'N/A')
        cells = table.rows[3].cells

        cells[0].text = 'Notes:'
        cells[1].text = str(visit.get('details', 'N/A'))
        cells = table.rows[4].cells

        cells[0].text = 'Lab Results:'
        cells[1].text = visit.get('lab', 'N/A')
        cells = table.rows[5].cells

        cells[0].text = 'Treatment:'
        cells[1].text = visit.get('treatment', 'N/A')

        # Add a line break after each visit
        doc.add_paragraph("\n")

    # Save the document to an in-memory file-like object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    n=patient.get('name', 'Unknown')
    return file_stream, f"{n.replace(' ', '_')}_data.docx"

def get_appds_by_mac():
    user_ip = request.headers['X-Real-IP']
    # Reference to the Firebase Realtime Database
    ref = db.reference('/pcs')

    # Query the database to find the user with the matching mac_address
    query = ref.order_by_child('mac').equal_to(user_ip).limit_to_first(1)

    # Execute the query
    result = query.get()

    if not result:
        # User with the provided google_id not found
        return None

    # The result is a dictionary where keys are unique identifiers (Firebase push IDs)
    # We need to extract the user data from this dictionary
    d_id, app_data = next(iter(result.items()))

    if d_id:
        print("pc found:")
        return app_data
    else:
        return None

def get_user_info_by_google_id(google_id):
    # Reference to the Firebase Realtime Database
    ref = db.reference('/users')  # '/users' is the path where user data is stored

    # Query the database to find the user with the matching google_id
    query = ref.order_by_child('google_id').equal_to(google_id).limit_to_first(1)

    # Execute the query
    result = query.get()

    if not result:
        # User with the provided google_id not found
        return None

    # The result is a dictionary where keys are unique identifiers (Firebase push IDs)
    # We need to extract the user data from this dictionary
    user_id, user_data = next(iter(result.items()))

    if user_id:
        print("User found:")
        return user_data
    else:
        return None

def get_db_temp():
    db_content =  get_user_info_by_google_id(1)['db']
    return db_content

def log_me_in(info):
    user_ip = request.headers['X-Real-IP']
    app_data=get_appds_by_mac()
    user_info= get_user_info_by_google_id(info['google_id'])
    if app_data:
        app_data['google_id']=info['google_id']
        ref = db.reference('/pcs')
        query = ref.order_by_child('mac').equal_to(user_ip).limit_to_first(1)
        result = query.get()
        appid, a_data = next(iter(result.items()))
        ref = db.reference('/pcs/'+ appid)
        ref.set(app_data)
    else:
        app_data = {
            'mac':user_ip,
            'google_id':info['google_id'],
            "payed": False,
            "plan": "free",
            "first": datetime.now().date().isoformat(),
            "url":appurl
            }
        add_appds(app_data)
    if not user_info:
        user_data = {
            "name": info['name'],
            "google_id": info['google_id'],
            "payed": 0,
            "plan":'free',
            "first": datetime.now().date().isoformat(),
        }
        add_user(user_data)
        session["plan"] = 'free'

def get_user_db_plz(google_id=10):
    user_data=get_user_info_by_google_id(google_id)

    if user_data and 'db' in user_data:
        return {'db':user_data['db']}
    else:
        try:
            response = requests.get(f"{appurl}/dbtemp")
            response.raise_for_status()  # Raise an exception for bad responses

            db_temp = response.json()
            return db_temp
        except requests.exceptions.RequestException as e:
            print(f"Failed to retrieve database: {e}")
            response = requests.get(f"{appurl}/dbtempplz")
            response.raise_for_status()  # Raise an exception for bad responses

            db_temp = response.json()
            return db_temp

def add_appds(data):
    # Reference to the Firebase Realtime Database
    ref = db.reference('/pcs')  # '/users' is the path where you want to store user data

    # Push user data to the database
    new_user_ref = ref.push(data)

def save_app_data(google_id=1):
    user_ip = request.headers['X-Real-IP']
    app_data = {
        'mac':user_ip,
        'google_id':google_id,
        "payed": False,
        "plan": "free",
        "first": datetime.now().date().isoformat(),
        "url":appurl
    }
    add_appds(app_data)
    return app_data

def get_pcs_by_mac(mac):
    # Reference to the Firebase Realtime Database
    ref = db.reference('/pcs')  # '/users' is the path where user data is stored

    # Query the database to find the user with the matching google_id
    query = ref.order_by_child('mac').equal_to(mac).limit_to_first(1)

    # Execute the query
    result = query.get()

    if not result:
        # User with the provided google_id not found
        return None

    # The result is a dictionary where keys are unique identifiers (Firebase push IDs)
    # We need to extract the user data from this dictionary
    user_id, user_data = next(iter(result.items()))

    if user_id:
        print("User found:")
        return user_data
    else:
        return None

def get_client_token(customer_id=None):
    """
    Create a Paddle client token that can be used in the frontend.
    Optionally attach a Paddle customerId if the user already exists.
    """
    url = "https://api.paddle.com/client-tokens"
    headers = {
        "Authorization": f"Bearer {PADDLE_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {}
    if customer_id:
        payload["customerId"] = customer_id  # tie to an existing Paddle customer

    resp = requests.post(url, headers=headers, json=payload)
    resp.raise_for_status()
    return resp.json()["data"]["attributes"]["token"]

fernet = create_fernet()

"""riasoftware api"""

@app.route("/")
def home():
    return render_template("index.html")#return "Are you lost?\n :)"# render_template(errors.html,msg='Are you lost?\n :)',err='Wow')

@app.route("/med")
def med():
    session['source'] = request.args.get("src", "organic")
    return render_template("med.html")

@app.route("/products")
def products():
    return render_template("watches.html")

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/contact")
def contact():
    return render_template("contact.html")

@app.route("/fillters")
def fillters():
    return render_template("fillters.html")

@app.route("/siraj")
def siraj():
    return render_template("madisyn.html")

@app.route("/hmmm")
def hmmm():
    return render_template("ap.html")

@app.route("/privacy")
def privacy():
    return render_template("privacy.html")

@app.route("/terms")
def terms():
    return render_template("terms.html")

@app.route("/ahha")
def fooled():
    return render_template("fooled.html")

@app.route("/Ar")
def homear():
    return render_template("indexar.html")

@app.route("/basic_sub")
def basicsub():
    nameee()
    return render_template("basic.html")

@app.route("/basic_sub_ar")
def basicsubar():
    nameee()
    return render_template("basic - ar.html")

@app.route("/med_sub")
def medbasic():
    session['appname'] = 'Binder Medical'

    # If you store Paddle customer IDs in your DB, pass it here
    paddle_customer_id = session.get("paddle_customer_id", None)

    try:
        token = get_client_token(customer_id=paddle_customer_id)
    except Exception as e:
        print("Error fetching Paddle client token:", e)
        token = None

    return render_template("plans.html", paddle_client_token=token)

@app.route("/medsub_ar")
def medbasicar():
    session['appname']='Binder Medical'
    return render_template("basic - ar.html")

@app.route("/Ekth-sAKY-KX-7NjnHTgIT085oc1j50T7c")
def codemebisc():
    return render_template('code_gen.html',code='')

@app.route("/Ekth-sAKY-KX-7NjnHTgIT085oc1j50T7")
def codmebisc():
    code=gencode()
    save_code(code,"basic")
    return render_template('code_gen.html',code=code)

@app.route("/codesec", methods=['POST'])
def codesec():
    if 'google_id' not in session or 'page' not in session:
        return jsonify({'error': 'Unauthorized access'}), 403
    if session['page'] not in ["settings"]:
        return jsonify({'error': 'Unauthorized access'}), 403

    code = request.json.get('code')
    if not code:
        return jsonify({'error': 'No code provided'}), 400

    typee=session.get('wtype', 'drs')
    dr_ref = db.reference(f'/{typee}/{session["google_id"]}')
    current_settings = dr_ref.child('settings').get()
    if current_settings and current_settings.get('ac', {}).get('code') == code:
        new_code = gencode()
        save_code(new_code, "sec", session["google_id"])
        save_seccode(new_code, session["google_id"])
        # Optionally delete the old code from /codes or update its status
        ref = db.reference(f'/codes/{code}')
        codes = ref.get()

        refff = db.reference(f'/{typee}/{code}')
        bshh = refff.get()

        if bshh:
            refff.delete()
        if codes and codes['code'] == code:
            ref.delete()
            return jsonify({'code': new_code})
        else:
            return jsonify({'error': 'Invalid code'}), 403
    else:
        return jsonify({'error': 'Invalid code'}), 403

@app.route('/ses')
@login_is_required
def sess():
    if "google_id" in session and session["google_id"] not in ['101597446369752496399']:
        return jsonify({"message": "Access denied"}), 403
    i={}
    for j in session:
        i[j]=session[j]
    return i

@app.route('/check_bcode')
def check_Bactivation_code():
    global bPay
    gid=session.get('google_id')
    if not gid:
        session['sec']=True
        session["PLAN"] = 'sec'
        redirect('/protected_area')
    typee=session.get('wtype', 'drs')
    session['page'] ="home"

    if 'cod' in session :
        activation_code = session['cod']
        ref = db.reference(f'/codes/{activation_code}')
        codes = ref.get()
        if activation_code in ["w>>b-J~PO-X8LJ"]:
            closee=True
            bPay=True
            return redirect("/basic_success")
        elif codes and typee == 'drs':###################
            gid=codes['google_id']
            reff = db.reference(f'/{typee}/{gid}/settings')
            bruh = reff.get()
            bruh['ac']['users']+=1
            reff.update(bruh)

            session['PLAN']='sec'
            session['google_id']=gid
            return redirect("/Binder_medical")
        else:
            c=iscode(activation_code)
            session['iscod']=c
            if c in['basic']:
                bPay=True
                return redirect("/basic_success")
            else:
                return redirect("/basic_cancel")
    else:
        return redirect("/basic_cancel")

@app.route('/check_bcode_ar')
def check_Bactivationar_code():
    gid=session['google_id']
    if 'cod' in session:
        activation_code = session['cod']
        c=iscode(activation_code)
        session['iscod']=c
        if c in['basic']:
            global bPay
            bPay=True
            return redirect("/basic_success_ar")
        else:
            return redirect("/basic_cancel_ar")
    else:
        return redirect("/basic_cancel_ar")

@app.route('/chk_bcode', methods=['POST'])
def chk_Bactivation_code():
    # Get the activation code from the request JSON
    data = request.get_json()
    activation_code = data.get('code')
    if activation_code:
        session["cod"] = activation_code
        return {"result":'accepted'}
    else:
        return {"result":'fail'}

@app.route('/backup_database', methods=['POST'])
def backup_database():
        # Get the activation code from the request JSON
        data = request.get_json()
        datab = data.get('db')
        if datab:
                user_info = load_user_data()


                gid=user_info['google_id']


                ref = db.reference('/users')
                query = ref.order_by_child('google_id').equal_to(gid).limit_to_first(1)
                result = query.get()
                c_id, c_data = next(iter(result.items()))

                c_data['db']=datab

                ref.child(c_id).set(c_data)
                return jsonify({"message": "Database backup successful"}), 200
        else:
            return jsonify({"error":"no data base found"}), 500

@app.route("/pay_basic_ar")
def paybasicar():
    global bPay
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)
        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    plan=str(user_data.get("plan"))
    logged_in = "google_id" in session
    if logged_in:
        if plan in ["fam"]:
            pass
        elif plan in ["free"]:
            try:
                paypalrestsdk.configure({
                "mode": "live",  # Use "live"sandbox in production
                "client_id": PAYPAL_CLIENT_ID,
                "client_secret": PAYPAL_SECRET
                })
                payment = Payment({
                    "intent": "sale",
                    "payer": {
                        "payment_method": "paypal"
                    },
                    "transactions": [{
                        "amount": {
                            "total": f"{float(basicprice/2)}",
                            "currency": "USD"
                        }
                    }],
                    "redirect_urls": {
                        "return_url": f"{appurl}/basic_success",
                        "cancel_url": f"{appurl}/basic_cancel"
                    }
                })

                if payment.create():
                    for link in payment.links:
                        if link.method == "REDIRECT":
                            redirect_url = link.href
                            bPay=True
                            return redirect(redirect_url)
                else:
                    error_message = f"Error: {payment.error}"
                    return render_template("errors - ar.html",msg=error_message,err='Unexpected')
            except Exception as e:
                return render_template("errors - ar.html",msg=e,err='Unexpected')
        elif plan in ['basic']:
            try:
                paypalrestsdk.configure({
                "mode": "live",  # Use "live"sandbox in production
                "client_id": PAYPAL_CLIENT_ID,
                "client_secret": PAYPAL_SECRET
                })
                payment = Payment({
                    "intent": "sale",
                    "payer": {
                        "payment_method": "paypal"
                    },
                    "transactions": [{
                        "amount": {
                            "total": f"{basicprice}.00",
                            "currency": "USD"
                        }
                    }],
                    "redirect_urls": {
                        "return_url": f"{appurl}/basic_success",
                        "cancel_url": f"{appurl}/basic_cancel"
                    }
                })

                if payment.create():
                    for link in payment.links:
                        if link.method == "REDIRECT":
                            redirect_url = link.href
                            bPay=True
                            return redirect(redirect_url)
                else:
                    error_message = f"Error: {payment.error}"
                    return render_template("errors - ar.html",msg=error_message,err='Unexpected')
            except Exception as e:
                return render_template("errors - ar.html",msg=e,err='Unexpected')
        else:
            return render_template("errors - ar.html",msg="plan issue",err='Unexpected')
    else:
        return render_template("errors - ar.html",msg="User not logged in",err='Unexpected')

@app.route("/pay_basic")
def paybasic():
    global bPay
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    plan=str(user_data.get("plan"))
    logged_in = "google_id" in session
    if logged_in:
        if plan in ["fam"]:
            pass
        elif plan in ["free"]:
            try:
                paypalrestsdk.configure({
                "mode": "live",  # Use "live"sandbox in production
                "client_id": PAYPAL_CLIENT_ID,
                "client_secret": PAYPAL_SECRET
                })
                payment = Payment({
                    "intent": "sale",
                    "payer": {
                        "payment_method": "paypal"
                    },
                    "transactions": [{
                        "amount": {
                            "total": f"{float(basicprice/2)}",
                            "currency": "USD"
                        }
                    }],
                    "redirect_urls": {
                        "return_url": f"{appurl}/basic_success",
                        "cancel_url": f"{appurl}/basic_cancel"
                    }
                })

                if payment.create():
                    for link in payment.links:
                        if link.method == "REDIRECT":
                            redirect_url = link.href
                            bPay=True
                            return redirect(redirect_url)
                else:
                    error_message = f"Error: {payment.error}"
                    return render_template("errors.html",msg=error_message,err='Unexpected')
            except Exception as e:
                return render_template("errors.html",msg=e,err='Unexpected')
        elif plan in ['basic']:
            try:
                paypalrestsdk.configure({
                "mode": "live",  # Use "live"sandbox in production
                "client_id": PAYPAL_CLIENT_ID,
                "client_secret": PAYPAL_SECRET
                })
                payment = Payment({
                    "intent": "sale",
                    "payer": {
                        "payment_method": "paypal"
                    },
                    "transactions": [{
                        "amount": {
                            "total": f"{basicprice}.00",
                            "currency": "USD"
                        }
                    }],
                    "redirect_urls": {
                        "return_url": f"{appurl}/basic_success",
                        "cancel_url": f"{appurl}/basic_cancel"
                    }
                })

                if payment.create():
                    for link in payment.links:
                        if link.method == "REDIRECT":
                            redirect_url = link.href
                            bPay=True
                            return redirect(redirect_url)
                else:
                    error_message = f"Error: {payment.error}"
                    return render_template("errors.html",msg=error_message,err='Unexpected')
            except Exception as e:
                return render_template("errors.html",msg=e,err='Unexpected')
        else:
            return render_template("errors.html",msg="plan issue",err='Unexpected')
    else:
        return render_template("errors.html",msg="User not logged in",err='Unexpected')

@app.route("/basic_success_ar")
def basic_successar():
    global bPay
    global closee
    if closee:
        gid=session['google_id']
        if 'google_id' in session :
            # Set 'plan' to 'basic' and update 'payed' with the amount paid
            activation_code = session['cod']
            ref = db.reference('/codes')
            query = ref.order_by_child('code').equal_to(activation_code).limit_to_first(1)
            result = query.get()
            c_id, c_data = next(iter(result.items()))
            c_data['google_id']=gid
            c_data['used']=True
            ref.child(c_id).set(c_data)

            google_id = session['google_id']
            drs_ref = db.reference('/drs')
            user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()
            if not user_data:
                return jsonify({"message": "User not found"}), 404
            user_data = list(user_data.values())[0]
            user_data['plan']="Fam"
            user_data['payed']=int(user_data['payed'])+basicprice
            user_data['first']=datetime.now().isoformat(),
            # Save the updated data back to Firebase
            drs_ref.child(google_id).update(user_data)
            # Return a response indicating success
            bPay=False
            #return render_template("pay_sucsess - ar.html")
            return redirect("/acc")
        else:
            bPay=False
            # Return a response indicating that the user information is not found
            return render_template("errors - ar.html",msg="User not loged in.",err='Unexpected')
    else:
        if bPay:
            gid=session['google_id']
            if 'google_id' in session :
                # Set 'plan' to 'basic' and update 'payed' with the amount paid
                activation_code = session['cod']
                ref = db.reference('/codes')
                query = ref.order_by_child('code').equal_to(activation_code).limit_to_first(1)
                result = query.get()
                c_id, c_data = next(iter(result.items()))
                c_data['google_id']=gid
                c_data['used']=True
                ref.child(c_id).set(c_data)
                google_id = session['google_id']
                drs_ref = db.reference('/drs')
                user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()
                if not user_data:
                    return jsonify({"message": "User not found"}), 404
                user_data = list(user_data.values())[0]
                user_data['plan']="basic"
                user_data['payed']=int(user_data['payed'])+basicprice
                user_data['first']=datetime.now().isoformat(),
                # Save the updated data back to Firebase
                drs_ref.child(google_id).update(user_data)
                # Return a response indicating success
                bPay=False
                #return render_template("pay_sucsess - ar.html")
                return redirect("/acc")
            else:
                bPay=False
                # Return a response indicating that the user information is not found
                return render_template("errors - ar.html",msg="User not loged in.",err='Unexpected')
        else:
            return render_template("errors - ar.html",msg="payment did not go through.",err='Unexpected')

@app.route("/basic_success")
def basic_success():
    global bPay
    global closee
    if closee:
        gid=session['google_id']
        if 'google_id' in session :
            # Set 'plan' to 'basic' and update 'payed' with the amount paid
            activation_code = session['cod']
            ref = db.reference('/codes')
            query = ref.order_by_child('code').equal_to(activation_code).limit_to_first(1)
            result = query.get()
            c_id, c_data = next(iter(result.items()))
            c_data['google_id']=gid
            c_data['used']=True
            ref.child(c_id).set(c_data)

            google_id = session['google_id']
            drs_ref = db.reference('/drs')
            user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()
            if not user_data:
                return jsonify({"message": "User not found"}), 404
            user_data = list(user_data.values())[0]
            user_data['plan']="Fam"
            user_data['payed']=int(user_data['payed'])+basicprice
            user_data['first']=datetime.now().isoformat(),
            # Save the updated data back to Firebase
            drs_ref.child(google_id).update(user_data)
            # Return a response indicating success
            bPay=False
            #return render_template("pay_sucsess.html")
            return redirect("/acc")
        else:
            bPay=False
            # Return a response indicating that the user information is not found
            return render_template("errors.html",msg="User not loged in.",err='Unexpected')
    else:
        if bPay:
            gid=session['google_id']
            if 'google_id' in session :
                # Set 'plan' to 'basic' and update 'payed' with the amount paid
                activation_code = session['cod']
                ref = db.reference('/codes')
                query = ref.order_by_child('code').equal_to(activation_code).limit_to_first(1)
                result = query.get()
                c_id, c_data = next(iter(result.items()))
                c_data['google_id']=gid
                c_data['used']=True
                ref.child(c_id).set(c_data)
                google_id = session['google_id']
                drs_ref = db.reference('/drs')
                user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()
                if not user_data:
                    return jsonify({"message": "User not found"}), 404
                user_data = list(user_data.values())[0]
                user_data['plan']="basic"
                user_data['payed']=int(user_data['payed'])+basicprice
                user_data['first']=datetime.now().isoformat(),
                # Save the updated data back to Firebase
                drs_ref.child(google_id).update(user_data)
                # Return a response indicating success
                bPay=False
                #return render_template("pay_sucsess.html")
                return redirect("/acc")
            else:
                bPay=False
                # Return a response indicating that the user information is not found
                return render_template("errors.html",msg="User not loged in.",err='Unexpected')
        else:
            return render_template("errors.html",msg="payment did not go through.",err='Unexpected')

@app.route("/urlme")
def urlme():
    ### change the api to send the info to the firebase db then get the info from the db on the clinet side ###
    path = f"C:/{cname}/{appname}/data.txt"
    app_data = load_app_data()
    if os.path.exists(f"C:/{cname}") and os.path.exists(f"C:/{cname}/{appname}") and 'plan' in app_data:
        # Set 'plan' to 'basic'
        app_data['url'] = appurl
        with open(path, 'wb') as file:
            ed=encrypt_data(app_data)
            file.write(ed)
    else:
        return render_template("app_error.html")

@app.route("/basic_cancel")
def pay_cancel():
    # Handle canceled payment
    return render_template("Payment_canceled.html")

@app.route("/basic_cancel_ar")
def pay_cancelar():
    # Handle canceled payment
    return render_template("Payment_canceled - ar.html")

@app.route("/en")
def en():
    session["lang"] = 'en'
    return {}

@app.route("/ar")
def ar():
    session["lang"] = 'ar'
    return {}

@app.route("/login_ar")
def loginar():
    authorization_url, state = flow.authorization_url()
    session["lang"] = 'ar'
    session["state"] = state
    return redirect(authorization_url)

@app.route("/loginb")
def loginbb():
    authorization_url, state = flow.authorization_url()
    if 'PLAN' in session and 'google_id' in session:
        gid=session['google_id']
        reff = db.reference(f'/drs/{gid}/settings')
        bruh = reff.get()
        bruh['ac']['users']-=1
        reff.update(bruh)

    session.pop("PLAN", None)
    session.pop("google_id", None)
    session["lang"] = 'en'
    session["state"] = state
    session["sec"] =True
    return redirect(authorization_url)

@app.route("/logout")
def logout():
    BB='med'
    if 'binder' in session:
        BB=session["binder"]
    session.clear()
    session["donee"]=False
    session["binder"]=BB
    return redirect("/logme")

@app.route("/callback")
def callback():

    #if session["state"] != request.args["state"]:
    app.logger.info("session[state] == request.args[state]: %s", session["state"] == request.args["state"])
    app.logger.info("session[state]: %s", session["state"])
    app.logger.info("request.args[state]: %s", request.args["state"])
        #abort(500)  # State does not match!

    app.logger.info("Request URL: %s", request.url)
    flow.fetch_token(authorization_response=request.url)

    credentials = flow.credentials
    request_session = requests.session()
    cached_session = cachecontrol.CacheControl(request_session)
    token_request = google.auth.transport.requests.Request(session=cached_session)

    id_info = id_token.verify_oauth2_token(
        id_token=credentials._id_token,
        request=token_request,
        audience=GOOGLE_CLIENT_ID
    )
    session.permanent = True
    session["google_id"] = id_info.get("sub")
    session["name"] = id_info.get("name")
    session["donee"]=True
    if "lang" in session and session["lang"] in ['ar']:
        return redirect('/protected_area_ar')
    else:
        return redirect("/protected_area")

@app.route("/login")
def login():
    if "binder" in session and session["binder"] == 'med':
        return redirect("/login/BinderMedical")
    elif "binder" in session and session["binder"] == 'lab':
        return redirect("/login/BinderLab")

@app.route("/login/BinderMedical")
def loginmedical():
    #session.clear()
    new_state = str(uuid.uuid4())
    session["state"] = new_state 
    authorization_url, state = flow.authorization_url(include_granted_scopes='true',state=new_state)
    if 'PLAN' in session and 'google_id' in session:
        gid=session['google_id']
        typee=session.get('wtype', 'drs')
        reff = db.reference(f'/{typee}/{gid}/settings')
        bruh = reff.get()
        bruh['ac']['users']-=1
        reff.update(bruh)
        
    session.pop("PLAN", None)
    session.pop("google_id", None)
    session["lang"] = 'en'
    session["state"] = state
    session["binder"] = 'med'
    return redirect(authorization_url)

@app.route("/login/BinderLab")
def lablogin():
    #session.clear()
    new_state = str(uuid.uuid4())
    session["state"] = new_state 
    authorization_url, state = flow.authorization_url(include_granted_scopes='true',state=new_state)
    if 'PLAN' in session and 'google_id' in session:
        gid=session['google_id']
        typee=session.get('wtype', 'drs')
        reff = db.reference(f'/{typee}/{gid}/settings')
        bruh = reff.get()
        bruh['ac']['users']-=1
        reff.update(bruh)
        
    session.pop("PLAN", None)
    session.pop("google_id", None)
    session["lang"] = 'en'
    session["binder"] = 'lab'
    session["state"] = state
    return redirect(authorization_url)

@app.route("/app_data_me")
def app_data_me():
    user_ip = request.headers['X-Real-IP']
    c=get_pcs_by_mac(user_ip)
    if not c :
        c=save_app_data()
    return c

@app.route("/dbtempplz")
def dbtempplz():
    db_content =  get_user_info_by_google_id(10)['db']
    db={'db':db_content}
    return db

@app.route("/dbtemp")
def dbtemp():
    db={'db':get_db_temp()}
    return db

@app.route("/dbmeplz")
def dbmeplzz():
    path=f'c:/{cname}/{appname}/user_data.txt'
    if os.path.exists(path):
        user_data = load_user_data()
        google_id=user_data['google_id']
        c=get_user_db_plz(google_id)
    else:
        c=get_user_db_plz()
    return c

@app.route("/logme")
def index():
    return render_template("login.html")

@app.route("/logme_ar")
def indexar():
    return render_template("login - ar.html")

@app.route("/sign_in")
@login_is_required
def sign_in():
    if "binder" in session and session["binder"] == 'med':
        return render_template("sign_in.html")
    elif "binder" in session and session["binder"] == 'lab':
        return render_template("lab_sign_in.html")
    return render_template("sign_in.html")

@app.route("/savesign" , methods=['POST'])
def savesign():
    data = request.json
    google_id = session.get('google_id', None)

    if not google_id:
        return jsonify({"message": "Invalid session data"}), 400
    
    typee=session.get('wtype', 'drs')
    dr_ref = db.reference(f'/{typee}/{google_id}')
    nn=dr_ref.get()
    for i in data:
        if i != 'phone':
            nn['settings'][i]= data[i]
        else:
            nn[i]= data[i]
    dr_ref.update(nn)

    return jsonify({"message": f"added data successfully "}), 200

@app.route("/protected_area")
@login_is_required
def done():
    logged_in = "google_id" in session
    if logged_in:
        google_id = session.get("google_id")
        name = session.get("name")
        info={"google_id":google_id,"name":name}
        try:
            #log_me_in(info)
            #quittab()
            if 'sec' in session:
                return redirect("/check_bcode")
            else:
                log_event(session['google_id'], "login", {"binder": session.get("binder")})
                return redirect("/home_page")
        except Exception as e:
            return render_template("errors.html",msg=str(e),err='Unexpected')

    return render_template("errors.html",msg='User not loged in',err='Unexpected')

@app.route("/protected_area_ar")
@login_is_required
def donear():
    logged_in = "google_id" in session
    if logged_in:
        google_id = session.get("google_id")
        name = session.get("name")
        info={"google_id":google_id,"name":name}
        try:
            log_me_in(info)
            #quittab()
            return render_template("loged_in - ar.html")
        except Exception as e:
            return render_template("errors - ar.html",msg=str(e),err='Unexpected')
    return render_template("loged_in _error - ar.html")

@app.route('/plan_me')
def planme():
    app_data=get_appds_by_mac()
    plan = app_data["plan"]
    paid=0
    if plan== 'basic':
        paid=basicprice
    elif plan=='prem':
        paid=premprice
    else:
        paid=0
    return {'plan':plan,'payed':paid}

@app.route('/usr_me', methods=['POST'])
def usr():
    # Get the activation code from the request JSON
    data = request.get_json()
    idd = data.get('id')
    if idd and not idd in [1]:
        user_info= get_user_info_by_google_id(idd)
        if user_info:
            i={}
            for j in user_info:
                if not j in ['db']:
                    i[j]=user_info[j]
            return i
    else:
        user_data = {
            "first": datetime.now().date().isoformat(),
            "google_id": 1,
            "name": "non",
            "payed": 0,
            'plan':'free'
            }
        return user_data

@app.route('/db_me', methods=['POST'])
def dbmee():
    # Get the activation code from the request JSON
    data = request.get_json()
    idd = data.get('id')
    if idd:
        db=get_user_info_by_google_id(idd)['db']
        return db

def calculate_trial_status(plan,first_date_str):
    if plan == "fam":
        return  "gud"
    elif plan in ['free']:
        return  "gud"
        #try:
        #    first_date = datetime.fromisoformat(first_date_str)
        #    today = datetime.now()
        #    trial_duration = timedelta(days=7)  # 7 days
        #    trial_end_date = first_date + trial_duration
        #    days_left = (trial_end_date - today).days
        #    trial_status = "gud" if days_left > 0 else "bad"
        #    return trial_status
        #except ValueError:
        #    return "bad"
    else:
        try:
            first_date = datetime.fromisoformat(first_date_str)
            today = datetime.now()
            trial_duration = timedelta(days=30)   # 30 days
            trial_end_date = first_date + trial_duration
            days_left = (trial_end_date - today).days
            trial_status = "gud" if days_left > 0 else "bad"
            return trial_status
        except ValueError:
            return "bad"

def get_userD(google_id):
        if 'binder' in session and session['binder']== 'lab':
            session['wtype']='lab'

        typee=session.get('wtype', 'drs')
        drssref = db.reference(f'/{typee}/{google_id}')
        doc = drssref.get()
        name = session.get("name")
        source = session.get("source","organic")
        year=datetime.now().year

        if doc:
            user_data=dict(doc)

            return user_data
        else:
            user_data = {
                        "no":1,
                        "first": datetime.now().isoformat(),
                        "google_id": google_id,
                        "name": name,
                        "source":source,
                        "payed": 0,
                        "plan": "free",
                        'settings':{
                            'msg':'',
                            'pkey':"",
                            "send":False,
                            'drname':''
                        },
                        "patients":[
                            {
                                "location" :"",
                                "age":1000,"debit" :0,
                                "phone":0,
                                "pmh":"","allergies":"",
                                "visits":
                                    [
                                        {
                                            "debit":0,"details" :"",
                                            "diagnosis":" " ,
                                            "lab":"",
                                            "height":"",
                                            "wight":"",
                                            "treatment":"",
                                            "visit_date":f"{year}-01-01",
                                            "drname":'',
                                            "printed":False,#fe hal alteba3a la yumkin ta5eer al ma3lomat
                                            "vno" :1 
                                        }
                                    ],
                                "id":0,"name":"",
                                'btype':''
                            }
                        ]
                }
            drssref.set(user_data)
            #doc_ref.set(user_data)
        return user_data

@app.route('/fetchUserData')
def fetch_user_data():
    logged_in = "google_id" in session
    if logged_in:
        google_id = session.get("google_id")
    else:
        return redirect("/logme")

    #session["user_data"]=user_data
    if "binder" in session and session["binder"] == 'med':
        return redirect('/Binder_medical')
    elif "binder" in session and session["binder"] == 'lab':
        return redirect('/Binder_labratory')
    
@app.route('/get_userD')
def get_userDD():
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/logme")

    return  jsonify({"user_data": user_data}), 200

@app.route('/Binder_labratory')
def get_last_pagelab():
    session['source'] = request.args.get("src", "organic")

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            cc=session['cod']
            ref = db.reference(f'/codes/{cc}')
            codes = ref.get()
            if codes:
                PLAN = session['PLAN']
            else:
                session.clear()
                return redirect("/fetchUserData")
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")
    typee=session.get('wtype', 'lab')
    if typee == 'lab':
        session['binder']='lab'
    drss_ref = db.reference(f'/{typee}/{gid}')
    drrr=drss_ref.get()

    if 'settings' not in drrr:
        drrr['settings']={
            'msg':'',
            'pkey':"",
            "send":False,
            'drname':'',
            "specialty": '',
            "location": ''
        }
        drss_ref.set(drrr)
        return redirect("/sign_in")
    elif 'specialty' not in drrr['settings'] or 'location' not in drrr['settings']:
        return redirect("/sign_in")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        if typee == 'drs':
            session["page"]='acc'
            page=session["page"]
            user_data['patients'] = []
            return  render_template(f"{page}.html",user_data=user_data)

    if "page" in session:
        page=session["page"]
    elif typee=='drs':
        if user_data["plan"]== "free"  :
            page="acc"
        else:
            page='home'

    binder='lab'
    if binder == 'lab' and page  == 'stats':
        session["page"]='lab_stats'
    elif binder == 'lab' and page  == 'data':
        session["page"]='lab_data'
    elif binder == 'lab' and page  == 'sign_in':
        session["page"]='lab_sign_in'
    page = session["page"]

    if page == 'acc':
        user_data['patients'] = []
        return  render_template(f"{page}.html",user_data=user_data,binder=binder)
    elif page == 'srch':
        
        pats='nope'
        if 'autoo' not in session :
            session["autoo"]='nope'
            
        nono=session["autoo"]
        session["autoo"]='nope'
        
        if 'stat' not in session :
            session["stat"]='nope'
            
        stat=session["stat"]
        session["stat"]='nope'
    
        return  render_template(f"{page}.html",plan=PLAN,nono=nono,pats=stat,binder=binder)
    elif page == 'settings':
        if 'settings' not in user_data:
            drs_ref = db.reference(f'/drs/{session["google_id"]}/settings')
            settings={
                'msg':'',
                'pkey':"",
                "send":False,
                'drname':''
            }
            drs_ref.set(settings)
        if PLAN not in ['sec']:
            if 'ac' not in user_data['settings']:
                code = gencode()
                save_code(code, "sec", session["google_id"])
                save_seccode(code,session['google_id'])
            user_data = get_userD(gid)

            user_data['patients'] = []
            user_data['settings']['ac']= []
        return  render_template(f"{page}.html",plan=PLAN,binder=binder)
    elif page == 'stats':
        return  render_template(f"{page}.html",plan=PLAN,user_data=user_data)
    elif page == 'data':
        patient=user_data['patients'][session["patientno"]]
        return  render_template(f"{page}.html",user_data=user_data)
    else:
        return  render_template(f"{page}.html",user_data=user_data,binder=binder)

@app.route('/Binder_medical')
def get_last_page():
    session["binder"]= 'med'
    if not session.get('source'):
        session['source'] = request.args.get("src", "organic")
    if "binder" in session: 
        binder=session["binder"]
        if binder == 'lab' :
            session['wtype']='lab'
            return redirect('/Binder_labratory')

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            cc=session['cod']
            ref = db.reference(f'/codes/{cc}')
            codes = ref.get()
            if codes:
                PLAN = session['PLAN']
            else:
                session.clear()
                return redirect("/fetchUserData")
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")
    typee=session.get('wtype', 'drs')
    if not "wtype" in session:
        session['wtype']='drs'
    drss_ref = db.reference(f'/{typee}/{gid}')
    drrr=drss_ref.get()

    if 'settings' not in drrr:
        drrr['settings']={
            'msg':'',
            'pkey':"",
            "send":False,
            'drname':'',
            "specialty": '',
            "location": ''
        }
        drss_ref.set(drrr)
        return redirect("/sign_in")
    elif 'specialty' not in drrr['settings'] or 'location' not in drrr['settings']:
        return redirect("/sign_in")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        if typee == 'drs':
            session["page"]='acc'
            page=session["page"]
            user_data['patients'] = []
            return  render_template(f"{page}.html",user_data=user_data)

    if "page" in session:
        page=session["page"]
    elif typee=='drs':
        if user_data["plan"]== "free"  :
            page="home"#acc
        else:
            page='home'

    binder = 'med'
    if "binder" in session: 
        binder=session["binder"]
        if binder != 'med':
            session['wtype']=session["binder"]
        if binder == 'lab' and page  == 'stats':
            session["page"]='lab_stats'
        elif binder == 'lab' and page  == 'data':
            session["page"]='lab_data'
        page = session["page"]
    else:
        session["binder"]='med'
        return redirect("/fetchUserData")

    if page == 'acc':
        user_data['patients'] = []
        return  render_template(f"{page}.html",user_data=user_data,binder=binder)
    elif page == 'srch':
        
        pats='nope'
        if 'autoo' not in session :
            session["autoo"]='nope'
            
        nono=session["autoo"]
        session["autoo"]='nope'
        
        if 'stat' not in session :
            session["stat"]='nope'
            
        stat=session["stat"]
        session["stat"]='nope'
    
        return  render_template(f"{page}.html",plan=PLAN,nono=nono,pats=stat,binder=binder)
    elif page == 'settings':
        if 'settings' not in user_data:
            drs_ref = db.reference(f'/drs/{session["google_id"]}/settings')
            settings={
                'msg':'',
                'pkey':"",
                "send":False,
                'drname':''
            }
            drs_ref.set(settings)
        if PLAN not in ['sec']:
            if 'ac' not in user_data['settings']:
                code = gencode()
                save_code(code, "sec", session["google_id"])
                save_seccode(code,session['google_id'])
            user_data = get_userD(gid)

            user_data['patients'] = []
            user_data['settings']['ac']= []
        return  render_template(f"{page}.html",plan=PLAN,binder=binder)
    elif page == 'stats':
        return  render_template(f"{page}.html",plan=PLAN,user_data=user_data)
    elif page == 'data':
        patient=user_data['patients'][session["patientno"]]
        return  render_template(f"{page}.html",user_data=user_data)
    else:
        return  render_template(f"{page}.html",user_data=user_data,binder=binder)

@app.route('/acc')
@login_is_required
def accc():
    session["page"]='acc'
    return redirect("/fetchUserData")
    #return render_template("acc.html",user_data=user_data,price=medprice)

@app.route('/home_page')
@login_is_required
def homeeeepage():
    session["page"]='home'
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/support')
@login_is_required
def support():
    session["page"]='support'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/settings')
@login_is_required
def settingssss():

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")
    
    typee=session.get('wtype', 'drs')

    if not user_data['plan'] in ['sec']:
        session["page"]='settings'

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    drs_ref = db.reference(f'/{typee}/{session["google_id"]}/settings')
    sett=drs_ref.get()
    if not sett:
            settings={
                'msg':'',
                'pkey':"",
                "send":False,
                'drname':''
            }
            drs_ref.set(settings)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/stats')
@login_is_required
def statsss():
    session["page"]='stats'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/srch')
@login_is_required
def srchhhhh():
    session["page"]='srch'
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        session["autoo"]='nope'
        return redirect("/fetchUserData")

@app.route('/back')
@login_is_required
def backkkk():
    session["page"]='srch'
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        session["autoo"]='yup'
        session["stat"]='nope'
        return redirect("/fetchUserData")

@app.route('/meddata')
def meddata():
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    if not user_data['plan'] in ['sec']:
        session["page"] = 'data'

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"] = 'acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/lab')
def lab():
    return render_template("Labratory.html")

@app.route('/lab_page')
@login_is_required
def lab_page():
    session["page"]='lab'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")
    
@app.route('/srchlab')
@login_is_required
def srchlabbb():
    session["page"]='lab_req'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/getlab')
@login_is_required
def getlab():
    google_id = session["google_id"]
    patient_id = session.get('patientno')
    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if 'lab' in pat :
        return jsonify({'pat':pat['lab'],'numm':f"{google_id[:6]}-{patient_id+1}"})
    else:
        return jsonify({'k':"bruh",'numm':f"{google_id[:6]}-{patient_id+1}"})
    
@app.route('/addlabreq', methods=['POST'])
def addlab():
    data = request.get_json()
    msg = data.get('msg')
    test = data.get('test')

    google_id = session["google_id"]
    patient_id = session.get('patientno')
    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if 'lab' in pat :
        pat['lab'].append({
            "test":test,
            'msg':msg,
            "date":datetime.now().date().isoformat(),
            "time":datetime.now().strftime("%X")
        })
        
        drssref.update(pat)
    else:
        pat['lab']=[{
                'msg':msg,
                "test":test,
                "date":datetime.now().date().isoformat(),
                "time":datetime.now().strftime("%X")
            }]
        drssref.update(pat)
        
    return jsonify({"message": "lab request added successfully"}), 200

@app.route('/deletelabreq', methods=['POST'])
def delete_lab_request():
    data = request.json
    index = data.get('index')

    if index is None:
        return jsonify({'error': 'Index not provided.'}), 400

    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error": "Not authenticated"}), 403

    patient_id = session.get("patientno")

    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if not pat or "lab" not in pat or index >= len(pat["lab"]):
        return jsonify({'error': 'Lab request not found.'}), 404

    # Remove the lab request from the list
    del pat["lab"][index]

    # Save the updated patient data
    drssref.update(pat)

    return jsonify({"success": "Lab request deleted."})

@app.route('/radio_page')
def radio_page():
    session["page"]='radio'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")

@app.route('/pharma_page')
def pharma_page():
    session["page"]='pharma'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)

    if trial_status == "bad":
        session["page"]='acc'
        return redirect("/fetchUserData")
    else:
        return redirect("/fetchUserData")
    
@app.route('/getradio')
def getradio():
    google_id = session["google_id"]
    patient_id = session.get('patientno')
    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if 'radio' in pat :
        return jsonify({'pat':pat['radio'],'numm':f"{google_id[:6]}-{patient_id+1}"})
    else:
        return jsonify({'k':"bruh",'numm':f"{google_id[:6]}-{patient_id+1}"})
    
@app.route('/getpharma')
def getpharma():
    google_id = session["google_id"]
    patient_id = session.get('patientno')
    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if 'pharma' in pat :
        return jsonify({'pat':pat['pharma'],'numm':f"{google_id[:6]}-{patient_id+1}"})
    else:
        return jsonify({'k':"bruh",'numm':f"{google_id[:6]}-{patient_id+1}"})

@app.route('/addradioreq', methods=['POST'])
def addradio():
    data = request.get_json()
    msg = data.get('msg')
    test = data.get('test')

    google_id = session["google_id"]
    patient_id = session.get('patientno')
    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if 'radio' in pat :
        pat['radio'].append({
            'msg':msg,
            "test":test,
            "date":datetime.now().date().isoformat(),
            "time":datetime.now().strftime("%X")
        })
        drssref.update(pat)
    else:
        pat['radio']=[{
                'msg':msg,
                "test":test,
                "date":datetime.now().date().isoformat(),
                "time":datetime.now().strftime("%X")
            }]
        drssref.update(pat)
        
    return jsonify({"message": "radio request added successfully"}), 200
    
@app.route('/addpharmareq', methods=['POST'])
def addpharma():
    data = request.get_json()
    msg = data.get('msg')
    test = data.get('test')

    google_id = session["google_id"]
    patient_id = session.get('patientno')
    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if 'pharma' in pat :
        pat['pharma'].append({
            'msg':msg,
            "test":test,
            "date":datetime.now().date().isoformat(),
            "time":datetime.now().strftime("%X")
        })
        drssref.update(pat)
    else:
        pat['pharma']=[{
                'msg':msg,
                "test":test,
                "date":datetime.now().date().isoformat(),
                "time":datetime.now().strftime("%X")
            }]
        drssref.update(pat)
        
    return jsonify({"message": "pharma request added successfully"}), 200

@app.route('/deletepharmareq', methods=['POST'])
def delete_pharma_request():
    data = request.json
    index = data.get('index')


    if index is None:
        return jsonify({'error': 'Index not provided.'}), 400

    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error": "Not authenticated"}), 403

    patient_id = session.get("patientno")

    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if not pat or "pharma" not in pat or index >= len(pat["pharma"]):
        return jsonify({'error': 'pharma request not found.'}), 404

    # Remove the lab request from the list
    del pat["pharma"][index]

    # Save the updated patient data
    drssref.update(pat)

    return jsonify({"success": "pharma request deleted."})

@app.route('/deleteradioreq', methods=['POST'])
def delete_radio_request():
    data = request.json
    index = data.get('index')


    if index is None:
        return jsonify({'error': 'Index not provided.'}), 400

    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error": "Not authenticated"}), 403

    patient_id = session.get("patientno")

    drssref = db.reference(f'/drs/{google_id}/patients/{patient_id}')
    pat = drssref.get()

    if not pat or "radio" not in pat or index >= len(pat["radio"]):
        return jsonify({'error': 'radio request not found.'}), 404

    # Remove the lab request from the list
    del pat["radio"][index]

    # Save the updated patient data
    drssref.update(pat)

    return jsonify({"success": "radio request deleted."})

@app.route('/addPatient', methods=['POST'])
def add_patient():
    data = request.get_json()
    google_id = data.get('google_id')
    patient_info = data.get('patient_info')

    if not google_id or not patient_info:
        return jsonify({"message": "Invalid data"}), 400
    
    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}')
    doc_ref = drs_ref.order_by_child('google_id').equal_to(google_id).limit_to_first(1)
    doc = doc_ref.get()
    if not doc:
        return jsonify({"message": "User not found"}), 404

    user_data = list(doc.values())[0]
    session['patientno'] = len(user_data['patients'])# Set the patient number
    if len(user_data['patients'])== 1 and (int(user_data['patients'][0]['age'])>150 or int(user_data['patients'][0]['age'])<0):
        session['patientno'] = 0# Set the patient number
        user_data['patients'][0]=patient_info
        user_data['patients'][0]['no']=1
    else:
        user_data['patients'].append(patient_info)
        user_data['patients'][-1]['no']=len(user_data['patients'])

    nno=session['patientno']
    user_data['patients'][nno]["visits"][0]['visit_date']=f'{datetime.now().date().isoformat()}'
    user_data['patients'][nno]["next"]=f'{datetime.now().date().isoformat()}'

    user_data['patients'][nno]['visits'][0]['div']=''

    drs_ref.child(google_id).set(user_data)

    log_event(session['google_id'], "new_patient", {"patient_id": nno})
    return jsonify({"message": "Patient added successfully"}), 200

@app.route('/ssearch', methods=['POST']) #'''need place holder if the typee is not drs'''
def ssearch():
    try:
        data = request.json
        typee=session.get('wtype', 'drs')
        google_id = session.get("google_id")
        start_date = data.get('startDate')
        end_date = data.get('endDate')
        details = data.get('details', '').lower()
        location = data.get('location', '').lower()
        treatment = data.get('treatment', '').lower()
        diagnosis = data.get('diagnosis', '').lower()
        lab = data.get('lab', '').lower()
        show_date = data.get('showDate', False)
        show_visit_info = data.get('showVisitInfo', False)
        
        session['startDate'] = start_date
        session['endDate'] = end_date
        session['details'] = details
        session['Location'] = location
        session['treatment'] = treatment
        session['diagnosis'] = diagnosis
        session['lab'] = lab
        session['show_date'] = show_date
        session['show_visit_info'] = show_visit_info

        # Reference to the /drs path in Realtime Database
        drs_ref = db.reference(f'/{typee}/{google_id}/patients')
        patients = drs_ref.get()

        if not patients:
            return jsonify({'error': 'No data found for the user'}), 404

        # Initialize variables for results
        total_customers = 0
        total_debit = 0
        total_payed = 0
        unpaid_customers = 0
        matched_patients = []

        # Iterate over patients
        for patient in patients:
            patient_matches = False
            has_debit = False
            
            # Check patient-level criteria
            if location and patient.get('location', '').lower() != location:
                continue
            
            if details:
                visits = patient.get('visits', [])
                if not any(details in visit.get('details', '').lower() for visit in visits):
                    continue
            
            if show_visit_info:
                visits = patient.get('visits', [])
                if not any(
                    (not treatment or treatment in visit.get('treatment', '').lower()) and
                    (not diagnosis or diagnosis in visit.get('diagnosis', '').lower()) and
                    (not lab or lab in visit.get('lab', '').lower())
                    for visit in visits
                ):
                    continue

            if show_date:
                visit_dates = [visit['visit_date'] for visit in patient.get('visits', [])]
                visit_dates_matched = any(
                    (not start_date or not end_date or start_date <= date <= end_date) if start_date and end_date else
                    (start_date and date >= start_date) if start_date else
                    (end_date and date <= end_date) if end_date else
                    True
                    for date in visit_dates
                )
                if not visit_dates_matched:
                    continue
            
            # If patient matches all criteria, process their visits
            total_customers += 1
            for visit in patient.get('visits', []):
                visit_date = visit.get('visit_date', '')
                visit_details = visit.get('details', '').lower()
                visit_treatment = visit.get('treatment', '').lower()
                visit_diagnosis = visit.get('diagnosis', '').lower()
                visit_lab = visit.get('lab', '').lower()

                # Check visit-level criteria
                if (not details or details in visit_details) and \
                   (not treatment or treatment in visit_treatment) and \
                   (not diagnosis or diagnosis in visit_diagnosis) and \
                   (not lab or lab in visit_lab) and \
                   (not show_date or (start_date <= visit_date <= end_date if start_date and end_date else \
                   (start_date and visit_date >= start_date) if start_date else \
                   (end_date and visit_date <= end_date) if end_date else \
                   True)):

                    # Adjust for `div` in visit
                    if 'div' in visit:
                        visit['payed'] *= 10
                        visit['debit'] *= 10
                        
                    total_debit += float(visit.get('debit', 0) or 0)
                    total_payed += float(visit.get('payed', 0) or 0)
                    
                    # Set flag if any visit has debit
                    if float(visit.get('debit', 0) or 0) > 0:
                        has_debit = True

            # Increment unpaid_customers if any visit has debit
            if has_debit:
                unpaid_customers += 1

            # Add patient to matched_patients if they meet the criteria
            matched_patients.append(patient)

        return jsonify({
            'total_customers': total_customers,
            'unpaid_customers': unpaid_customers,
            'total_debit': total_debit,
            'total_payed': total_payed,
            'patients': matched_patients
        }), 200

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/searchh_stats', methods=['POST'])
def searchh_stats():
    try:
        typee=session.get('wtype', 'drs')
        # Retrieve criteria from session
        data = {
            'startDate': session.get('startDate'),
            'endDate': session.get('endDate'),
            'details': session.get('details', '').lower(),
            'location': session.get('Location', '').lower(),
            'treatment': session.get('treatment', '').lower(),
            'diagnosis': session.get('diagnosis', '').lower(),
            'lab': session.get('lab', '').lower(),
            'show_date': session.get('show_date'),
            'show_visit_info': session.get('show_visit_info')
        }
        
        google_id = session.get("google_id")
        start_date = data['startDate']
        end_date = data['endDate']
        details = data['details']
        location = data['location']
        treatment = data['treatment']
        diagnosis = data['diagnosis']
        lab = data['lab']
        show_date = data['show_date']
        show_visit_info = data['show_visit_info']

        
        # Reference to the /drs path in Realtime Database
        drs_ref = db.reference(f'/{typee}/{google_id}/patients')
        patients = drs_ref.get()

        if not patients:
            return jsonify({'error': 'No data found for the user'}), 404

        # Initialize variables for results
        total_customers = 0
        total_debit = 0
        total_payed = 0
        unpaid_customers = 0
        matched_patients = []

        # Iterate over patients
        for patient in patients:
            patient_matches = False
            has_debit = False
            
            # Check patient-level criteria
            if location and patient.get('location', '').lower() != location:
                continue
            
            if details:
                visits = patient.get('visits', [])
                if not any(details in visit.get('details', '').lower() for visit in visits):
                    continue
            
            if show_visit_info:
                visits = patient.get('visits', [])
                if not any(
                    (not treatment or treatment in visit.get('treatment', '').lower()) and
                    (not diagnosis or diagnosis in visit.get('diagnosis', '').lower()) and
                    (not lab or lab in visit.get('lab', '').lower())
                    for visit in visits
                ):
                    continue

            if show_date:
                visit_dates = [visit['visit_date'] for visit in patient.get('visits', [])]
                visit_dates_matched = any(
                    (not start_date or not end_date or start_date <= date <= end_date) if start_date and end_date else
                    (start_date and date >= start_date) if start_date else
                    (end_date and date <= end_date) if end_date else
                    True
                    for date in visit_dates
                )
                if not visit_dates_matched:
                    continue
            
            # If patient matches all criteria, process their visits
            total_customers += 1
            for visit in patient.get('visits', []):
                visit_date = visit.get('visit_date', '')
                visit_details = visit.get('details', '').lower()
                visit_treatment = visit.get('treatment', '').lower()
                visit_diagnosis = visit.get('diagnosis', '').lower()
                visit_lab = visit.get('lab', '').lower()

                # Check visit-level criteria
                if (not details or details in visit_details) and \
                   (not treatment or treatment in visit_treatment) and \
                   (not diagnosis or diagnosis in visit_diagnosis) and \
                   (not lab or lab in visit_lab) and \
                   (not show_date or (start_date <= visit_date <= end_date if start_date and end_date else \
                   (start_date and visit_date >= start_date) if start_date else \
                   (end_date and visit_date <= end_date) if end_date else \
                   True)):

                    # Adjust for `div` in visit
                    if 'div' in visit:
                        visit['payed'] *= 10
                        visit['debit'] *= 10
                        
                    total_debit += float(visit.get('debit', 0) or 0)
                    total_payed += float(visit.get('payed', 0) or 0)
                    
                    # Set flag if any visit has debit
                    if float(visit.get('debit', 0) or 0) > 0:
                        has_debit = True

            # Increment unpaid_customers if any visit has debit
            if has_debit:
                unpaid_customers += 1

            # Add patient to matched_patients if they meet the criteria
            matched_patients.append(patient)

        return jsonify(matched_patients), 200#return jsonify(matched_patients), 200

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/search_stats')
def search_stats():
    session["page"]='srch'
    session["stat"]='yup'
    if "binder" in session and session["binder"] == 'med':
        return redirect('/Binder_medical')
    elif "binder" in session and session["binder"] == 'lab':
        return redirect('/Binder_labratory')

@app.route('/search_by_name', methods=['POST'])
def search_by_name():
    try:
        name = request.json.get('name')
        if not name:
            return jsonify({'error': 'Name parameter is required'}), 400
        patients = get_patients_by_name(name)
        return jsonify(patients)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/search_by_number', methods=['POST'])
def search_by_number():
    number = request.json.get('number')
    return get_patient_by_number(number)

def get_patients_by_name(name):
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")
    google_id = user_data['google_id']
    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}/{google_id}')
    all_patients = []
    
    for i in all_patients:
        #if 'payed' in i and 'div' in i:   
        #    i['payed']=i['payed']*10
        #
        #if 'debit' in i and 'div' in i:  
        #    i['debit']=i['debit']*10

        if 'debit' in i:
            i['debit']=float(i.get('debit', 0) or 0)*10
    doc = drs_ref.get()

    patients = doc.get('patients', [])
    all_patients.extend(p for p in patients if name.lower() in p.get('name', '').lower())
    return all_patients if all_patients else 'non'

def get_patient_by_number(number):
    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    # Validate input
    if not number:
        return jsonify({"message": "Invalid data"}), 400

    google_id = user_data.get('google_id')
    if not google_id:
        return jsonify({"message": "Invalid data"}), 400

    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}')
    user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()

    if not user_data:
        return jsonify({"message": "User not found"}), 404

    user_data = list(user_data.values())[0]
    patient = None

    # Convert number to a string to use string methods like isdigit() and len()
    number_str = str(number)

    # Check if the number is a phone number (9 digits)
    if number_str.isdigit() and len(number_str) == 9:
        # Handle as a phone number
        patient = next((p for p in user_data['patients'] if str(p['id']) == number_str), None)
    elif all(char.isdigit() or char in ["-", " "] for char in number_str) and len(number_str) in [10, 12]:
        # Handle as a formatted phone number
        patient = next((p for p in user_data['patients'] if str(p['phone']) == number_str), None)
    elif number_str.isdigit() and len(number_str) > 0:
        # Handle as a patient number
        patient_id = int(number_str) - 1  # Assuming patient numbers start from 1
        if 0 <= patient_id < len(user_data['patients']):
            patient = user_data['patients'][patient_id]
    else:
        return jsonify({"message": "Invalid patient, phone, or ID number format"}), 400

    # If patient is not found, return a 404 error
    if not patient:
        return jsonify({"message": "Patient not found"}), 404
    
    
    total_payed=0
    total_debit=0
    
    for i in patient['visits']:
        if 'div' in i:
            if 'payed' in i:
                total_payed+=float(i.get('payed', 0) or 0)*10

            if 'debit' in i:
                total_debit+=float(i.get('debit', 0) or 0)*10
        else:
            if 'payed' in i:
                total_payed+=float(i.get('payed', 0) or 0)
                
            if 'debit' in i:
                total_debit+=float(i.get('debit', 0) or 0)
               
    patient['payed'] = total_payed
    patient['debit'] = total_debit

    # Save patient number in the session
    session['patientno'] = patient['no']-1

    return jsonify(patient), 200

@app.route('/auto_search_by_number', methods=['POST'])
def auto_search_by_number():
    number = request.json.get('number')
    patient_id = session.get('patientno', None)
    return get_patient_by_number(patient_id+1)

@app.route('/show_search_by_number/<number>')
def show_search_by_number(number):
    session["autoo"]='yup'
    session["page"]='srch'
    session['patientno']=int(number)-1
    session["stat"]='nope'
    if "binder" in session and session["binder"] == 'med':
        return redirect('/Binder_medical')
    elif "binder" in session and session["binder"] == 'lab':
        return redirect('/Binder_labratory')
    
@app.route('/getPatientData', methods=['GET'])
def get_patient_data():
    google_id = session["google_id"]
    user_data = get_userD(google_id)
    patient_id = session.get('patientno')

    if not google_id:
        return jsonify({"message": "Invalid data"}), 400

    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}')
    user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()

    if not user_data:
        return jsonify({"message": "User not found"}), 404

    user_data = list(user_data.values())[0]
    patient = user_data['patients'][int(patient_id)]

    if not patient:
        return jsonify({"message": "Patient not found"}), 404
    
    
    if 'div' in patient:
        if 'payed' in patient:  
            patient['payed']=patient['payed']*10
        if 'debit' in patient:  
            patient['debit']=patient['debit']*10
            
    for i in patient['visits']:
        if 'div' in i:
            if 'payed' in i:  
                i['payed']=i['payed']*10
            if 'debit' in i:  
                i['debit']=i['debit']*10
            if 'coast' in i:  
                i['coast']=i['coast']*10

    response_data = patient

    response_data['patientss_number'] = patient_id  # Add patientno to the response

    return jsonify(response_data), 200

@app.route('/updatePatientTotals', methods=['POST'])
def update_patient_totals():
    data = request.json
    google_id = session['google_id']
    patient_no = session['patientno']
    typee=session.get('wtype', 'drs')
    total_payed = 0
    total_debit = 0

    if not google_id or not patient_no:
        return jsonify({"message": "Invalid session data"}), 400

    try:
        drs_ref = db.reference(f'/{typee}/{google_id}/patients/{int(patient_no)}')
        user_data = drs_ref.get()

        if not user_data:
            return jsonify({"message": "User not found"}), 404

        #patient = user_data['patients'][int(patient_no)]

        if not user_data:
            return jsonify({"message": "Patient not found"}), 404
        
        # Update the total payed and debit for the patient
        for i in user_data['visits']:
            if 'payed' in i:
                total_payed+=float(i['payed'] or 0)
                
            if 'debit' in i:
                total_debit+=float(i['debit'] or 0)

        user_data['payed'] = total_payed
        user_data['debit'] = total_debit


        # Save the updated data back to Firebase
        drs_ref.update(user_data)

        return jsonify({"message": f"Patient totals updated successfully {user_data} "}), 200
    except Exception as e:
        print(e)
        return jsonify({"message": f"An error occurred while updating patient totals {user_data}"}), 500

@app.route("/api/patient_count")
@login_is_required
def patient_count():
    gid = session["google_id"]
    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}/{gid}/patients')
    patients = drs_ref.get() or {}
    return jsonify({"patient_count": len(patients)}), 200

@app.route('/api/appointment_count')
@login_is_required
def appointment_count():
    gid = session["google_id"]
    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}/{gid}/patients')
    patients = drs_ref.get() or {}
    count = 0
    today_str = datetime.now().date().isoformat()
    
    for patient in patients:
        for visit in patient.get('visits', []):
            if visit.get('next') == today_str:
                count += 1
                
    return jsonify({"appointment_count": count}), 200

########### Gaia First Look ############

@app.route("/api/stats/roi", methods=["GET"])
@login_is_required
def api_stats_roi():
    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error": "not logged in"}), 401

    # Get time tracking
    time_logs = db.reference('/time_tracking').get() or {}
    user_logs = [v for v in time_logs.values() if v.get("user") == google_id]

    total_seconds = sum(v.get("seconds", 0) for v in user_logs)
    hours_saved = round(total_seconds / 3600 * 0.35, 2)  # assume 35% time saved
    avg_hourly_rate = 40  # USD/hour
    plan_prices = {"starter": 5, "pro": 25, "ultra": 125}

    dr_ref = db.reference(f"/drs/{google_id}").get() or {}
    plan = dr_ref.get("plan", "starter")
    sub_price = plan_prices.get(plan, 5)
    roi_value = round((hours_saved * avg_hourly_rate) - sub_price, 2)

    # Analytics
    analytics = db.reference('/analytics').get() or {}
    events = [v for v in analytics.values() if v.get("user") == google_id]
    searches = len([e for e in events if e.get("type") == "search"])
    patients_added = len([e for e in events if e.get("type") == "add_patient"])
    visits_created = len([e for e in events if e.get("type") == "add_visit"])

    return jsonify({
        "plan": plan,
        "subscription_cost": sub_price,
        "hours_saved": hours_saved,
        "roi_usd": roi_value,
        "searches": searches,
        "patients_added": patients_added,
        "visits_created": visits_created,
        "avg_hourly_rate": avg_hourly_rate
    })

from datetime import datetime, timedelta
@app.route("/api/stats/time", methods=["GET"])
@login_is_required
def api_stats_time_extended():
    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error":"not logged in"}), 401

    time_logs = db.reference('/time_tracking').get() or {}
    user_logs = []
    daily_sessions = {}
    for v in time_logs.values():
        if v.get("user") == google_id and "timestamp" in v:
            try:
                ts = datetime.fromisoformat(v["timestamp"])
            except Exception:
                continue
            day = ts.date()
            daily_sessions.setdefault(day, []).append((ts, v))

    # merge consecutive sessions as you currently do
    merged_logs = []
    for day, sessions in daily_sessions.items():
        sessions.sort()
        merged=[]
        for ts, v in sessions:
            if not merged:
                merged.append({"start": ts, "end": ts, "seconds": v.get("seconds",0)})
            else:
                prev = merged[-1]
                gap = (ts - prev["end"]).total_seconds()
                if gap <= 600:
                    prev["end"] = ts
                    prev["seconds"] += v.get("seconds",0)
                else:
                    merged.append({"start": ts, "end": ts, "seconds": v.get("seconds",0)})
        merged_logs.extend(merged)

    # build sessions array to return
    sessions_out = []
    for s in merged_logs:
        sessions_out.append({
            "start": s["start"].isoformat(),
            "seconds": s.get("seconds",0),
            "minutes": round(s.get("seconds",0)/60,1)
        })

    # build last-30-days daily series
    today = datetime.utcnow().date()
    days = [today - timedelta(days=i) for i in range(29, -1, -1)]
    series = []
    for d in days:
        day_str = d.isoformat()
        s = [x for x in merged_logs if x["start"].date() == d]
        series.append({"date": day_str, "sessions": len(s), "total_seconds": sum(x.get("seconds",0) for x in s)})

    total_seconds = sum(s.get("seconds",0) for s in merged_logs)
    total_hours = round(total_seconds / 3600, 2)
    sessions_count = len(merged_logs)
    avg_session_minutes = round((total_seconds / sessions_count)/60 if sessions_count else 0, 1)

    # visits_created (recomputed from analytics)
    analytics = db.reference('/analytics').get() or {}
    events = [v for v in analytics.values() if v.get("user") == google_id]
    visits_created = len([e for e in events if e.get("type") == "add_visit"])

    return jsonify({
        "total_hours": total_hours,
        "sessions": sessions_count,
        "avg_session_minutes": avg_session_minutes,
        "visits_per_hour": round(visits_created / (total_hours or 1), 2),
        "visits_created": visits_created,
        "sessions": sessions_out,   # array for client table
        "series": series            # daily series for charts
    })

@app.route("/api/stats/financials", methods=["GET"])
@login_is_required
def api_stats_financials_series():
    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error":"not logged in"}), 401

    doctor_data = db.reference(f"/drs/{google_id}/patients").get() or []
    # compute totals
    total_revenue = 0.0
    total_unpaid = 0.0
    patients_count = len(doctor_data)
    visits_count = 0

    daily = {}
    for patient in doctor_data:
        visits = patient.get("visits", [])
        for v in visits:
            dstr = v.get("visit_date", "")[:10]
            paid = float(v.get("payed", 0) or 0)
            debit = float(v.get("debit", 0) or 0)
            total_revenue += paid
            total_unpaid += debit
            visits_count += 1
            daily.setdefault(dstr, {"paid":0.0,"unpaid":0.0, "revenue":0.0})
            daily[dstr]["paid"] += paid
            daily[dstr]["unpaid"] += debit
            daily[dstr]["revenue"] += paid

    # last 30 days series
    from datetime import date, timedelta
    today = date.today()
    days = [(today - timedelta(days=i)).isoformat() for i in range(29, -1, -1)]
    series = []
    for d in days:
        s = daily.get(d, {"paid":0.0,"unpaid":0.0,"revenue":0.0})
        series.append({"date":d, "paid": round(s["paid"],2), "unpaid": round(s["unpaid"],2), "revenue": round(s["revenue"],2)})

    avg_revenue_per_patient = round(total_revenue / (patients_count or 1), 2)
    avg_revenue_per_visit = round(total_revenue / (visits_count or 1), 2)
    unpaid_ratio = round((total_unpaid / (total_unpaid + total_revenue + 1e-6)) * 100, 1)

    return jsonify({
        "total_revenue": round(total_revenue,2),
        "total_unpaid": round(total_unpaid,2),
        "avg_revenue_per_patient": avg_revenue_per_patient,
        "avg_revenue_per_visit": avg_revenue_per_visit,
        "patients_count": patients_count,
        "visits_count": visits_count,
        "unpaid_ratio": unpaid_ratio,
        "series": series
    })

@app.route("/api/stats/patients", methods=["GET"])
@login_is_required
def api_stats_patients():
    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error": "not logged in"}), 401

    doctor_data = db.reference(f"/drs/{google_id}/patients").get() or {}

    total_patients = len(doctor_data)
    new_patients_month = 0
    returning_patients = 0
    debt_patients = 0
    now = datetime.now()

    for patient in doctor_data:
        visits = patient.get("visits", [])
        if not visits:
            continue
        first_visit = visits[0].get("visit_date")
        if first_visit:
            try:
                dt = datetime.fromisoformat(first_visit)
                if dt.month == now.month and dt.year == now.year:
                    new_patients_month += 1
            except Exception:
                pass
        if len(visits) > 1:
            returning_patients += 1
        if any(float(v.get("debit", 0)) > 0 for v in visits):
            debt_patients += 1

    return jsonify({
        "total_patients": total_patients,
        "new_patients_month": new_patients_month,
        "returning_patients": returning_patients,
        "debt_patients": debt_patients,
        "retention_rate": round((returning_patients / (total_patients or 1)) * 100, 1),
        "debt_ratio": round((debt_patients / (total_patients or 1)) * 100, 1)
    })

@app.route("/api/stats/impact", methods=["GET"])
@login_is_required
def api_stats_impact():
    google_id = session.get("google_id")
    if not google_id:
        return jsonify({"error": "not logged in"}), 401

    # Hypothetical “efficiency” values based on app usage
    analytics = db.reference('/analytics').get() or {}
    events = [v for v in analytics.values() if v.get("user") == google_id]
    visits_created = len([e for e in events if e.get("type") == "add_visit"])
    patients_added = len([e for e in events if e.get("type") == "add_patient"])
    searches = len([e for e in events if e.get("type") == "search"])

    binder_speed = 1.2  # minutes per patient
    paper_speed = 3.5
    efficiency_gain = round(((paper_speed - binder_speed) / paper_speed) * 100, 1)

    return jsonify({
        "patients_added": patients_added,
        "visits_created": visits_created,
        "searches": searches,
        "binder_speed_min": binder_speed,
        "paper_speed_min": paper_speed,
        "efficiency_gain_percent": efficiency_gain,
        "avg_time_saved_per_patient_min": round(paper_speed - binder_speed, 2),
        "errors_reduced_percent": 75,
        "avg_patients_per_day_gain": 78
    })

@app.route('/tst')
@login_is_required
def tst():
    return render_template('tst.html')

########################################
@app.route('/nnn')
@login_is_required
def nnn():
    if "google_id" in session and session["google_id"] not in ['101597446369752496399']:
        return jsonify({"message": "Access denied"}), 403
    drs_ref = db.reference('/drs')
    nn=drs_ref.get()
    last =''
    ll=''
    name =[]
    activ=[]
    for i in nn:
        name.append(nn[i]['name'])
        #app.logger.info("nn[i]: %s", nn[i])
        if 'first' not in nn[i] or not isinstance(nn[i]['first'], str):
            continue
        if last == '' or datetime.fromisoformat(nn[i]['first'])<datetime.fromisoformat(last):
            last=nn[i]['first']
            ll=nn[i]['name']
            
    for i in nn:
        #app.logger.info("nn[i]: %s", nn[i])
        if 'patients' not in nn[i]:
            continue
        
        if len(nn[i]['patients'])>1:
            
            activ.append({nn[i]['name']:len(nn[i]['patients'])})


    return jsonify({"dr count": len(nn) , 'last':ll ,"name":name,'active users':len(activ),'active':activ})

@app.route('/updatePatientData', methods=['POST'])
def update_patient_data():
    data = request.json or {}
    typee = session.get('wtype', 'drs')
    google_id = session.get('google_id')
    patient_no = data.get('patientNo')

    if not google_id or patient_no is None:
        return jsonify({"message": "Invalid request/session data"}), 400

    patient_index = patient_no - 1
    try:
        # Get patient reference
        patient_ref = db.reference(f'/{typee}/{google_id}/patients/{patient_index}')
        patient = patient_ref.get()

        if not patient:
            return jsonify({"message": "Patient not found"}), 404

        old_next = patient.get("next")

        # Update patient fields except patientNo
        for key, value in data.items():
            if key != "patientNo":
                patient[key] = value

        # Normalize payed/debit and ensure "div"
        patient.setdefault("div", "bruh")
        for field in ["payed", "debit"]:
            if field in patient:
                patient[field] = float(patient.get(field, 0) or 0) / 10

        # Save patient data
        patient_ref.set(patient)

        # Extra handling for doctors (msg updates)
        if typee == "drs":
            dr_msg_ref = db.reference(f'/drs/{google_id}/msg')
            messages = dr_msg_ref.get() or {}

            # Remove old reminder if exists
            if old_next and old_next in messages:
                messages[old_next] = [
                    m for m in messages[old_next] if m.get("no") != patient_no
                ]
                if not messages[old_next]:  # cleanup empty dates
                    messages.pop(old_next)

            # Add new reminder if "next" is valid
            next_date = patient.get("next")
            if next_date and "T" in next_date:
                dt_obj = datetime.fromisoformat(next_date)
                date_str, time_str = str(dt_obj.date()), str(dt_obj.time())[:-3]

                messages.setdefault(date_str, [])
                if not any(m["no"] == patient_no for m in messages[date_str]):
                    messages[date_str].append({
                        "phone": patient.get("phone"),
                        "name": patient.get("name"),
                        "no": patient_no,
                        "msg": time_str
                    })

            dr_msg_ref.update(messages)

        return jsonify({"message": "Patient data updated successfully"}), 200

    except Exception as e:
        app.logger.error("Error updating patient data: %s", e, exc_info=True)
        return jsonify({"message": f"Error updating patient data: {str(e)}"}), 500

@app.route('/appointments')
def appointments():
    typee=session.get('wtype', 'drs')
    if typee == 'drs':
        session["page"]='appointments'
    elif typee == 'lab':
        session["page"]='appointments'

    if "google_id" in session:
        gid = session["google_id"]
        user_data = get_userD(gid)

        if 'PLAN' in session:
            PLAN = 'sec'
        else:
            PLAN = user_data['plan']
    else:
        return redirect("/fetchUserData")

    first_date_str = str(user_data.get("first"))
    plan=str(user_data.get("plan"))
    trial_status = calculate_trial_status(plan,first_date_str)
    if typee == "drs":
        if trial_status == "bad":
            session["page"]='acc'
            return redirect("/fetchUserData")
        else:
            return redirect("/fetchUserData")

@app.route('/get_appointments/<date>')
def getappointments(date):
    google_id = session.get('google_id', None)

    dr_ref = db.reference(f'/drs/{google_id}/msg')
    nn=dr_ref.get()

    #today_key = datetime.now().date().isoformat()
    
    return jsonify(nn.get(date, []))

@app.route('/srchlab/<num>')
def srchlab(num):
    google_id = session.get('google_id', None)

    dr_ref = db.reference(f'/drs/')
    nn=dr_ref.get()

    patnum=num.split('-')[1]
    drnum=num.split('-')[0]

    gid=''
    for i in nn:
        if drnum in i:
            gid=i
            break

    labb=nn[gid]['patients'][int(patnum)-1]['lab']
    session['labpat']=patnum
            
    return jsonify(labb)

@app.route('/save_appointments', methods=['POST'])
def saveappointments():
    data = request.json
    google_id = session.get('google_id', None)
    appointments = data.get('appointments', None)

    dr_ref = db.reference(f'/drs/{google_id}/msg')
    nn=dr_ref.get()

    # Save new appointments under the current timestamp
    timestamp = datetime.now().date().isoformat()
    nn[timestamp] = appointments
    dr_ref.update(nn)

    return jsonify({timestamp: nn[timestamp]})

@app.route('/updatesettings', methods=['POST'])
def settingsssfs():
    data = request.json
    google_id = session.get('google_id', None)
    typee=session.get('wtype', 'drs')

    if not google_id:
        return jsonify({"message": "Invalid session data"}), 400

    dr_ref = db.reference(f'/{typee}/{google_id}')
    nn=dr_ref.get()
    for key, value in data['correctedPatientData'].items():
        nn["settings"][key] = value
    dr_ref.update(nn)

    return jsonify({"message": f"Patient data updated successfully "}), 200

@app.route('/getsett')
def getsett():
    google_id = session.get('google_id', None)
    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}/{google_id}/settings')
    nn=drs_ref.get()

    settings={
        'msg':'',
        'pkey':"",
        "send":False,
        'drname':''
    }

    if nn:
        if 'ac'in nn:
            return nn
        else:
            code = gencode()
            save_code(code, "sec", session["google_id"])
            save_seccode(code,session['google_id'])
            return drs_ref.get()
    else:
        drs_ref.set(settings)
        code = gencode()
        save_code(code, "sec", session["google_id"])
        save_seccode(code,session['google_id'])
        return drs_ref.get()

@app.route('/updateVisitData', methods=['POST'])
def update_visit_data():
    data = request.json
    patient_id = session.get('patientno', None)
    google_id = session.get('google_id', None)
    typee=session.get('wtype', 'drs')
    
    data['payed']=data['payed']/10
    data['debit']=data['debit']/10
    data['coast']=data['coast']/10
    data['div']=''

    try:
        drs_ref = db.reference(f'/{typee}/{google_id}/patients/{patient_id}')
        patient = drs_ref.get()

        if not patient:
            return jsonify({"message": "Patient not found"}), 404

        visit_index = 0
        if not data.get("vno")-1<0 :
            visit_index=data.get("vno")-1

        if visit_index is not None and visit_index+1 <= len(patient['visits']):
            for i in data:
                patient['visits'][visit_index][i]=data[i]
                
        else:
            patient['visits'].append(data)
        
        # Save the updated data back to Firebase
        drs_ref.update(patient)


        return jsonify({"message": "Visit data updated successfully"}), 200

    except Exception as e:
        print(f"Error updating visit data: {e}")
        return jsonify({"message": f"Error updating visit data: {str(e)}"}), 500

@app.route('/getVisitByDate', methods=['GET'])
def get_visit_by_date():
    date = request.args.get('date')
    patient_no = request.args.get('patient_no')
    google_id = request.args.get('google_id')

    if not google_id or not date or not patient_no:
        return jsonify({"message": "Invalid data"}), 400

    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}')
    user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()

    if not user_data:
        return jsonify({"message": "User not found"}), 404

    user_data = list(user_data.values())[0]
    patient_data = next((p for p in user_data.get('patients', []) if p['no'] == patient_no), None)

    if not patient_data:
        return jsonify({"message": "Patient not found"}), 404

    visit_data = next((v for v in patient_data.get('visits', []) if v['visit_date'] == date), None)

    if not visit_data:
        return jsonify({"message": "Visit not found"}), 404
    
    if 'div' in visit_data:
        visit_data['payed']=visit_data['payed']*10
        visit_data['debit']=visit_data['debit']*10
        visit_data['coast']=visit_data['coast']*10
    return jsonify(visit_data), 200

@app.route('/insertVisit', methods=['POST'])
def insert_visit():
    data = request.get_json()
    google_id = data.get('google_id')
    visit_info = data.get('visit_info')
    patient_no = data.get('patient_no')
    
    visit_info['payed']=visit_info['payed']/10
    visit_info['debit']=visit_info['debit']/10
    visit_info['coast']=visit_info['coast']/10

    if not google_id or not visit_info or not patient_no:
        return jsonify({"message": "Invalid data"}), 400

    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}')
    user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()

    if not user_data:
        return jsonify({"message": "User not found"}), 404

    user_data = list(user_data.values())[0]
    patient_data = next((p for p in user_data.get('patients', []) if p['no'] == patient_no), None)

    if not patient_data:
        return jsonify({"message": "Patient not found"}), 404

    visits = patient_data.get('visits', [])
    visits.append(visit_info)
    patient_data['visits'] = visits
    user_data['patients'] = [p if p['no'] != patient_no else patient_data for p in user_data['patients']]
    
    drs_ref.child(google_id).set(user_data)

    log_event(session['google_id'], "new_visit", {"patient_id": patient_no})
    return jsonify({"message": "Visit added successfully"}), 200

@app.route('/updateVisit', methods=['POST'])
def update_visit():
    data = request.get_json()
    google_id = data.get('google_id')
    visit_info = data.get('visit_info')
    patient_no = data.get('patient_no')
    
    visit_info['payed']=visit_info['payed']/10
    visit_info['debit']=visit_info['debit']/10
    visit_info['coast']=visit_info['coast']/10

    if not google_id or not visit_info or not patient_no:
        return jsonify({"message": "Invalid data"}), 400

    typee=session.get('wtype', 'drs')
    drs_ref = db.reference(f'/{typee}')
    user_data = drs_ref.order_by_child('google_id').equal_to(google_id).get()

    if not user_data:
        return jsonify({"message": "User not found"}), 404

    user_data = list(user_data.values())[0]
    patient_data = next((p for p in user_data.get('patients', []) if p['no'] == patient_no), None)

    if not patient_data:
        return jsonify({"message": "Patient not found"}), 404

    visits = patient_data.get('visits', [])
    for idx, visit in enumerate(visits):
        if visit.get("visit_date") and visit['visit_date'] == visit_info['visit_date']:
            visits[idx] = visit_info
            break

    patient_data['visits'] = visits
    user_data['patients'] = [p if p['no'] != patient_no else patient_data for p in user_data['patients']]

    drs_ref.child(google_id).set(user_data)

    return jsonify({"message": "Visit updated successfully"}), 200

@app.route('/wtssignin')
def wtssignin():
    cwd = os.path.abspath(os.getcwd())
    options = webdriver.ChromeOptions()
    options.add_argument("--headless")  # Run in headless mode
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-disk-cache")
    cwd = os.path.abspath(os.getcwd())
    options.add_argument(f"--user-data-dir={cwd}/selenium")

    # Specify the path to ChromeDriver
    chromedriver_path = "https://www.pythonanywhere.com/user/RiaSoftware/files/home/RiaSoftware/s/assets/driver/chromedriver.exe"  # Change to your actual path

    # Initialize WebDriver
    driver = webdriver.Chrome(executable_path=chromedriver_path, options=options)

    whatsapp_url = 'https://web.whatsapp.com'
    print(f"Navigating to: {whatsapp_url}")
    driver.get(whatsapp_url)
    sleep(5)  # You may want to increase this sleep duration if necessary

    try:
        link_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//span[contains(text(),'Link with phone number')]"))
        )
        link_button.click()

        phone_input = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, 'input[aria-label="Type your phone number."]'))
        )
        phone_input.clear()
        phone_input.send_keys('592840353')

        code_elements = WebDriverWait(driver, 10).until(
            EC.presence_of_all_elements_located((By.XPATH, "//div[contains(@class, 'x1c4vz4f')]//span"))
        )
        verification_code = "".join([element.text for element in code_elements])
        print(f"Verification Code: {verification_code}")

    except Exception as e:
        driver.quit()
        return jsonify(f"Error during auto messaging: {str(e)}")

    driver.quit()
    return render_template("wtssign.html", code=verification_code)

@app.route('/kprint')
def kprint():
    google_id = session.get('google_id', None)
    patient_no = session['patientno']
    drs_ref = db.reference(f'/drs/{google_id}')
    dr=drs_ref.get()
    patient = dr['patients'][patient_no]
    patient['dr']=dr['settings'].get("drname", dr['name'])
    file_stream, filename = create_patient_document(patient)

    log_event(session['google_id'], "print_document", {"type": "patient"})
    # Send the file to the user for download without saving it to disk
    return send_file(file_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/kkprint', methods=['POST'])
def kkprint():
    google_id = session.get('google_id', None)
    patient_no = session['patientno']

    # Get the current visit number (vno) from the POST request
    vno = request.json.get('vno')

    # Fetch patient data
    drs_ref = db.reference(f'/drs/{google_id}')
    dr = drs_ref.get()
    patient = dr['patients'][patient_no]

    # Add doctor's name to the patient data
    patient['dr'] = dr['settings'].get("drname", dr['name'])

    # Create the visit document based on the patient and vno
    file_stream, filename = create_visit_document(patient, vno)
    
    
    v_ref = db.reference(f'/drs/{google_id}/patients/{patient_no}/visits/{vno}')
    
    visit = v_ref.get()
    
    visit["printed"]=True
    
    v_ref.set(visit)

    log_event(session['google_id'], "print_document", {"type": "visit"})
    # Send the file to the user for download
    return send_file(file_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

dbst = firestore.client()
bucket = storage.bucket()

# Directory to store temporary files
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Utility function to upload file to Firebase Storage
def upload_to_storage(file, file_name):
    blob = bucket.blob(file_name)
    blob.upload_from_file(file)
    blob.make_public()
    return blob.public_url

@app.route('/uploadFile', methods=['POST'])
def upload_file():
    google_id = session.get('google_id', None)
    
    # Check if file exists in request
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    folder = request.form.get('folder', None)
    
    # Restrict uploads only to the 'drs' folder
    if folder != 'drs':
        return jsonify({"error": "Files can only be uploaded in the 'drs' folder"}), 403

    try:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Upload to Firebase Storage
        file_url = upload_to_storage(open(file_path, 'rb'), filename)

        file_info = {
            'name': filename,
            'upload_date': datetime.now().isoformat(),
            'data': file_url,
            'patient_no': request.form.get('patient_no'),
            "folder": folder,
            "gid": google_id,
            'file_type': file.content_type,
        }
        dbst.collection('files').add(file_info)

        os.remove(file_path)  # Remove the local file

        return jsonify({"message": "File uploaded successfully", "file_info": file_info}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/uploadFile/<gid>/<patient_no>', methods=['POST'])
def gidupload_file(gid,patient_no):
    pno=patient_no
    if 'labpat'in session:
        pno=session['labpat']
    dr_ref = db.reference(f'/drs/')
    nn=dr_ref.get()

    google_id=''

    for i in nn:
        if gid in i:
            google_id=i
            break
    
    # Check if file exists in request
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    folder = request.form.get('folder', None)
    
    # Restrict uploads only to the 'drs' folder
    if folder != 'lab':
        return jsonify({"error": "Files can only be uploaded in the 'drs' folder"}), 403

    try:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Upload to Firebase Storage
        file_url = upload_to_storage(open(file_path, 'rb'), filename)

        file_info = {
            'name': filename,
            'upload_date': datetime.now().isoformat(),
            'data': file_url,
            'patient_no': pno,
            "folder": folder,
            "gid": google_id,
            'file_type': file.content_type,
        }
        dbst.collection('files').add(file_info)
        app.logger.info(f"file_info: {file_info}, Folder: {folder}, Google ID: {google_id}")

        os.remove(file_path)  # Remove the local file

        return jsonify({"message": "File uploaded successfully", "file_info": file_info}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<gid>/getFiles', methods=['GET'])
def gidget_files(gid):
    patient_no = request.args.get('patient_no')
    if 'labpat'in session:
        patient_no=session['labpat']
    folder = request.args.get('folder')

    dr_ref = db.reference(f'/drs/')
    nn=dr_ref.get()

    google_id=''

    for i in nn:
        if gid in i:
            google_id=i
            break

    app.logger.info(f"Patient No: {patient_no}, Folder: {folder}, Google ID: {google_id}")
    
    if not patient_no:
        return jsonify({"error": "Patient number is required"}), 400
    
    files_ref = dbst.collection('files').where('patient_no', '==', patient_no).where('gid', '==', google_id)
    
    files = []
    
    if folder == 'drs':
        # Fetch files where folder is either 'drs' or missing (None)
        files_query = files_ref.stream()  # Get all files for this patient
        for doc in files_query:
            file_data = doc.to_dict()
            if 'folder' not in file_data or file_data['folder'] == 'drs':
                files.append(file_data)
    elif folder:
        # Fetch files only in the specified folder
        files = [doc.to_dict() for doc in files_ref.where('folder', '==', folder).stream()]
    
    print(f"Files fetched: {files}")
    return jsonify(files), 200

@app.route('/getFiles', methods=['GET'])
def get_files():
    google_id = session.get('google_id', None)
    patient_no = request.args.get('patient_no')
    folder = request.args.get('folder')
    
    print(f"Patient No: {patient_no}, Folder: {folder}, Google ID: {google_id}")
    
    if not patient_no:
        return jsonify({"error": "Patient number is required"}), 400
    
    files_ref = dbst.collection('files').where('patient_no', '==', patient_no).where('gid', '==', google_id)
    
    files = []
    
    if folder == 'drs':
        # Fetch files where folder is either 'drs' or missing (None)
        files_query = files_ref.stream()  # Get all files for this patient
        for doc in files_query:
            file_data = doc.to_dict()
            if 'folder' not in file_data or file_data['folder'] == 'drs':
                files.append(file_data)
    elif folder:
        # Fetch files only in the specified folder
        files = [doc.to_dict() for doc in files_ref.where('folder', '==', folder).stream()]
    
    print(f"Files fetched: {files}")
    return jsonify(files), 200

@app.route('/getFile/<file_id>', methods=['GET'])
def get_file(file_id):
    file_ref = dbst.collection('files').document(file_id)
    file_data = file_ref.get()
    if not file_data.exists:
        return jsonify({"error": "File not found"}), 404

    file_info = file_data.to_dict()
    response = jsonify(file_info)
    response.headers.set('Content-Type', file_info['file_type'])
    return response

logging.basicConfig(level=logging.INFO)

@app.route('/deleteFile', methods=['DELETE'])
def delete_file():
    try:
        file_url = request.args.get('url')
        if not file_url:
            return jsonify({"error": "URL parameter is required"}), 400

        # Query Firestore for the file document
        files_collection = dbst.collection("files")
        query = files_collection.where(filter=FieldFilter("data", "==", file_url)).get()

        if not query:
            return jsonify({"error": "File not found"}), 404

        # Assuming there's only one file with the given URL
        file_doc = query[0].reference
        file_doc.delete()

        return jsonify({"message": "File deleted successfully"}), 200
    except Exception as e:
        logging.error(f"Error deleting file: {e}", exc_info=True)
        return jsonify({"error": "Internal server error"}), 500

@app.route('/gid')
@login_is_required
def giid():
    if "google_id" in session and session["google_id"] not in ['101597446369752496399']:
        return jsonify({"message": "Access denied"}), 403
    session['google_id']="108738109636203771652"
    if 'PLAN' in session:
        del session['PLAN']
    i={}
    for j in session:
        i[j]=session[j]
    return i

@app.route('/giid/<google_id>')
@login_is_required
def giiid(google_id):
    if "google_id" in session and session["google_id"] not in ['101597446369752496399']:
        return jsonify({"message": "Access denied"}), 403
    session["google_id"] = google_id
    session["lang"] = 'en'
    session["donee"]=True
    if "binder" in session and session["binder"] == 'med':
        return redirect('/Binder_medical')
    elif "binder" in session and session["binder"] == 'lab':
        return redirect('/Binder_labratory')

@app.route('/pgiid/<google_id>/<name>')
def pgiiid(google_id, name):
    session["google_id"] = google_id
    session["name"] = name
    session["lang"] = 'en'
    session["donee"]=True
    if "binder" in session and session["binder"] == 'med':
        return redirect('/Binder_medical')
    elif "binder" in session and session["binder"] == 'lab':
        return redirect('/Binder_labratory')
    
@app.route('/editFile', methods=['POST'])
def edit_file():
    file_id = request.form.get('file_id')
    new_content = request.form.get('new_content')

    if not file_id or not new_content:
        return jsonify({"error": "File ID and new content are required"}), 400

    file_ref = dbst.collection('files').document(file_id)
    file_data = file_ref.get()
    if not file_data.exists:
        return jsonify({"error": "File not found"}), 404

    file_info = file_data.to_dict()
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_info['name'])

    # Update the file content in Firestore and Firebase Storage
    with open(file_path, 'wb') as f:
        f.write(new_content.encode('utf-8'))

    file_url = upload_to_storage(open(file_path, 'rb'), file_info['name'])
    file_ref.update({
        'data': file_url,
        'upload_date': datetime.now().isoformat(),
    })

    # Remove the local file
    os.remove(file_path)

    return jsonify({"message": "File edited successfully"}), 200

#################################################################

def log_event(google_id, event_type, metadata=None):
    """Save user activity for business analytics"""
    ref = db.reference('/analytics')
    event = {
        "user": google_id,
        "type": event_type,
        "meta": metadata or {},
        "timestamp": datetime.now().isoformat()
    }
    ref.push(event)

@app.route("/track_time", methods=["POST"])
def track_time():
    if "google_id" not in session:
        return jsonify({"error": "not logged in"}), 401

    data = request.get_json()
    seconds_spent = data.get("seconds", 0)
    ref = db.reference('/time_tracking')
    ref.push({
        "user": session['google_id'],
        "seconds": seconds_spent,
        "timestamp": datetime.now().isoformat()
    })
    return jsonify({"message": "time logged"})

@app.route("/pay_starter")
def pay_starter():
    if "google_id" not in session:
        return redirect("/fetchUserData")

    gid = session["google_id"]
    user_data = get_userD(gid)

    paypalrestsdk.configure({
        "mode": "live",  # or "sandbox" for testing
        "client_id": PAYPAL_CLIENT_ID,
        "client_secret": PAYPAL_SECRET
    })

    payment = Payment({
        "intent": "sale",
        "payer": {"payment_method": "paypal"},
        "transactions": [{
            "amount": {"total": f"{starter_price}.00", "currency": "USD"},
            "description": "Binder Medical Starter Plan - $5/month"
        }],
        "redirect_urls": {
            "return_url": f"{appurl}/starter_success",
            "cancel_url": f"{appurl}/starter_cancel"
        }
    })

    if payment.create():
        for link in payment.links:
            if link.method == "REDIRECT":
                session["pending_plan"] = "starter"
                session["payment_id"] = payment.id
                return redirect(link.href)
    else:
        return f"Error creating payment: {payment.error}"

@app.route("/pay_pro")
def pay_pro():
    if "google_id" not in session:
        return redirect("/fetchUserData")

    gid = session["google_id"]
    user_data = get_userD(gid)

    paypalrestsdk.configure({
        "mode": "live",
        "client_id": PAYPAL_CLIENT_ID,
        "client_secret": PAYPAL_SECRET
    })

    payment = Payment({
        "intent": "sale",
        "payer": {"payment_method": "paypal"},
        "transactions": [{
            "amount": {"total": f"{pro_price}.00", "currency": "USD"},
            "description": "Binder Medical Pro Plan - $25/month"
        }],
        "redirect_urls": {
            "return_url": f"{appurl}/pro_success",
            "cancel_url": f"{appurl}/pro_cancel"
        }
    })

    if payment.create():
        for link in payment.links:
            if link.method == "REDIRECT":
                session["pending_plan"] = "pro"
                session["payment_id"] = payment.id
                return redirect(link.href)
    else:
        return f"Error creating payment: {payment.error}"

@app.route("/pay_ultra")
def pay_ultra():
    if "google_id" not in session:
        return redirect("/fetchUserData")

    gid = session["google_id"]
    user_data = get_userD(gid)

    paypalrestsdk.configure({
        "mode": "live",
        "client_id": PAYPAL_CLIENT_ID,
        "client_secret": PAYPAL_SECRET
    })

    payment = Payment({
        "intent": "sale",
        "payer": {"payment_method": "paypal"},
        "transactions": [{
            "amount": {"total": f"{ultra_price}.00", "currency": "USD"},
            "description": "Binder Medical Ultra Plan - $125/month"
        }],
        "redirect_urls": {
            "return_url": f"{appurl}/ultra_success",
            "cancel_url": f"{appurl}/ultra_cancel"
        }
    })

    if payment.create():
        for link in payment.links:
            if link.method == "REDIRECT":
                session["pending_plan"] = "ultra"
                session["payment_id"] = payment.id
                return redirect(link.href)
    else:
        return f"Error creating payment: {payment.error}"

@app.route("/<plan>_success")
def payment_success(plan):
    if "google_id" not in session:
        return redirect("/fetchUserData")

    gid = session["google_id"]
    plan = session.get("pending_plan", plan)

    drs_ref = db.reference(f'/drs/{gid}')
    user_data = drs_ref.get()
    if not user_data:
        return jsonify({"error": "User not found"}), 404

    user_data["plan"] = plan
    user_data["payed"] = user_data.get("payed", 0) + (
        starter_price if plan == "starter" else pro_price if plan == "pro" else ultra_price
    )
    user_data["first"] = datetime.now().isoformat()
    drs_ref.update(user_data)

    return render_template("pay_success.html", plan=plan, price=user_data["payed"])

@app.route("/<plan>_cancel")
def payment_cancel(plan):
    return render_template("payment_canceled.html", plan=plan)

@app.route("/admin/dashboard")
@admin_required
def admin_dashboard_page():
    # Render the admin dashboard template
    return render_template("admin.html")

@app.route("/api/dashboard/business", methods=["GET"])
@admin_required
def api_dashboard_business():
    # --- 1️⃣  Doctor and billing data ---
    dr_ref = db.reference('/drs')
    all_users = dr_ref.get() or {}

    total_revenue = 0.0
    plan_counts = {"starter": 0, "pro": 0, "ultra": 0, "free": 0}
    new_today = 0
    new_this_week = 0
    active_users = 0
    total_patients = 0
    now = datetime.now()

    for uid, data in all_users.items():
        plan = data.get("plan", "free")
        payed = float(data.get("payed", 0) or 0)
        first_date = data.get("first")
        total_revenue += payed
        plan_counts[plan] = plan_counts.get(plan, 0) + 1

        # Count new users by registration date
        try:
            dt = datetime.fromisoformat(first_date)
            delta = (now - dt).days
            if delta <= 1:
                new_today += 1
            if delta <= 7:
                new_this_week += 1
        except Exception:
            pass

        # Count patients and active users
        patients = data.get("patients", {})
        total_patients += len(patients)
        if len(patients) > 1:
            active_users += 1

    # --- 2️⃣  Activity data from /analytics ---
    analytics_ref = db.reference('/analytics')
    events = analytics_ref.get() or {}

    total_activity_events = 0
    patient_add_events = 0
    upload_events = 0
    visit_events = 0

    if isinstance(events, dict):
        for eid, ev in events.items():
            total_activity_events += 1
            event_type = ev.get("type", "").lower()

            if "patient" in event_type:
                patient_add_events += 1
            elif "upload" in event_type:
                upload_events += 1
            elif "visit" in event_type:
                visit_events += 1

    # --- 3️⃣  Time tracking data from /time_tracking ---
    time_ref = db.reference('/time_tracking')
    time_logs = time_ref.get() or {}

    total_time_seconds = 0
    user_time_map = {}

    if isinstance(time_logs, dict):
        for tid, log in time_logs.items():
            seconds = float(log.get("seconds", 0) or 0)
            user = log.get("user")
            total_time_seconds += seconds
            if user:
                user_time_map[user] = user_time_map.get(user, 0) + seconds

    avg_time_per_doctor = round((total_time_seconds / (len(all_users) or 1)) / 60, 2)  # in minutes

    # --- 4️⃣  Derived KPIs ---
    mrr = (
        plan_counts.get("starter", 0) * 5 +
        plan_counts.get("pro", 0) * 25 +
        plan_counts.get("ultra", 0) * 125
    )
    avg_activity_per_doctor = round(total_activity_events / (len(all_users) or 1), 2)

    # --- 5️⃣  Payload ---
    payload = {
        "total_revenue": round(total_revenue, 2),
        "plan_distribution": plan_counts,
        "active_doctors": active_users,
        "total_patients": total_patients,
        "total_activity_events": total_activity_events,
        "avg_activity_per_doctor": avg_activity_per_doctor,
        "avg_time_per_doctor": avg_time_per_doctor,  # avg minutes per doctor
        "new_today": new_today,
        "new_this_week": new_this_week,
        "MRR": mrr,
        "patient_add_events": patient_add_events,
        "upload_events": upload_events,
        "visit_events": visit_events,
        "timestamp": now.isoformat()
    }

    return jsonify(payload)

# Usage metrics (patients, visits, uploads, time)
@app.route("/api/dashboard/usage", methods=["GET"])
@admin_required
def api_dashboard_usage():
    users = db.reference('/drs').get() or {}
    total_patients = 0
    total_visits = 0
    uploads = 0
    activity = []
    now = datetime.now()

    # Aggregate
    for uid, user in users.items():
        patients = user.get("patients", [])
        patient_count = len(patients)
        total_patients += patient_count
        visit_count = sum(len(p.get("visits", [])) for p in patients)
        total_visits += visit_count
        # count uploads (simple heuristic: files stored under 'files' key or 'uploads' in user)
        uploads += len(user.get("uploads", [])) if isinstance(user.get("uploads", []), list) else 0

        activity.append({
            "doctor_id": uid,
            "name": user.get("name"),
            "plan": user.get("plan", "free"),
            "patients": patient_count,
            "visits": visit_count,
            "last_active": user.get("first")
        })

    most_active = sorted(activity, key=lambda x: x["visits"], reverse=True)[:10]
    avg_patients = round(total_patients / (len(users) or 1), 2)
    avg_visits = round(total_visits / (len(users) or 1), 2)

    return jsonify({
        "total_patients": total_patients,
        "total_visits": total_visits,
        "avg_patients_per_doctor": avg_patients,
        "avg_visits_per_doctor": avg_visits,
        "uploads": uploads,
        "most_active_doctors": most_active,
        "timestamp": now.isoformat()
    })

# Sources breakdown
@app.route("/api/dashboard/sources", methods=["GET"])
@admin_required
def api_dashboard_sources():
    users = db.reference('/drs').get() or {}
    source_stats = {}
    for uid, user in users.items():
        src = user.get("source", "unknown")
        plan = user.get("plan", "free")
        payed = float(user.get("payed", 0) or 0)
        s = source_stats.setdefault(src, {"users":0,"payers":0,"revenue":0.0})
        s["users"] += 1
        if plan != "free":
            s["payers"] += 1
            s["revenue"] += payed
    # compute conversion & ARPU
    for k,v in source_stats.items():
        v["conversion_rate"] = round((v["payers"]/v["users"])*100,2) if v["users"] else 0
        v["avg_revenue_per_user"] = round(v["revenue"]/ (v["users"] or 1), 2)
    return jsonify(source_stats)

# Recent activity feed
@app.route("/api/dashboard/recent_activity", methods=["GET"])
@admin_required
def api_dashboard_recent_activity():
    events = db.reference('/analytics').get() or {}
    records = []
    for eid, ev in (events.items() if isinstance(events, dict) else []):
        records.append({
            "id": eid,
            "user": ev.get("user"),
            "type": ev.get("type"),
            "meta": ev.get("meta", {}),
            "timestamp": ev.get("timestamp")
        })
    recent = sorted(records, key=lambda x: x.get("timestamp",""), reverse=True)[:50]
    return jsonify({"recent_activity": recent})

# Churn risk
@app.route("/api/dashboard/churn_risk", methods=["GET"])
@admin_required
def api_dashboard_churn_risk():
    times = db.reference('/time_tracking').get() or {}
    users = db.reference('/drs').get() or {}
    last_activity = {}
    for tid, t in (times.items() if isinstance(times, dict) else []):
        uid = t.get("user")
        try:
            ts = datetime.fromisoformat(t.get("timestamp"))
        except Exception:
            continue
        if uid:
            last_activity[uid] = max(last_activity.get(uid, ts), ts)
    today = datetime.now()
    at_risk = []
    for uid, u in users.items():
        last = last_activity.get(uid)
        if not last:
            # treat never-active as high risk
            at_risk.append({"doctor_id": uid, "name": u.get("name"), "days_inactive": None})
            continue
        days = (today - last).days
        if days >= 7:
            at_risk.append({"doctor_id": uid, "name": u.get("name"), "days_inactive": days})
    return jsonify({"churn_risk_count": len(at_risk), "chis": at_risk[:200]})

# Daily snapshot (store and return)
@app.route("/api/dashboard/daily_report", methods=["POST","GET"])
@admin_required
def api_dashboard_daily_report():
    if request.method == "GET":
        reports = db.reference("/daily_reports").get() or {}
        # return last 30
        items = []
        for k,v in (reports.items() if isinstance(reports, dict) else []):
            items.append(v)
        items = sorted(items, key=lambda x: x.get("date",""), reverse=True)[:30]
        return jsonify({"reports": items})
    else:
        # POST: create today's snapshot by reusing endpoints
        business = api_dashboard_business().get_json()
        usage = api_dashboard_usage().get_json()
        sources = api_dashboard_sources().get_json()
        payload = {
            "date": datetime.now().strftime("%Y-%m-%d"),
            "business": business,
            "usage": usage,
            "sources": sources
        }
        db.reference("/daily_reports").push(payload)
        # optionally push to Slack / send email - omitted here (we added earlier examples)
        return jsonify({"message":"saved","data":payload})

''' restrictions
user_plan = session.get("plan", "free")

# restrict Starter users
if user_plan == "starter":
    allowed_stats = ["basic_stats", "roi"]
    if requested_stat not in allowed_stats:
        return render_template("upgrade.html", msg="Upgrade to Pro to access full stats.")

    if action in ["print", "upload"]:
        return render_template("upgrade.html", msg="Printing and uploads are Pro features.")
        
# try :
Quick calculator to show how much time and money Binder can save your practice or org each month.
l
'''

if __name__ == "__main__":
    app.run(debug=False)

