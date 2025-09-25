"""
knowlage_base/pinecone.py

Pinecone-backed persistence with:
 - DocumentConverter strategies (CSV, TXT, PDF, DOCX, JSON, Generic)
 - LoaderFactory to pick converter by file extension
 - ToDocument context that coordinates conversion + chunking
 - PineconeStore accepts an 'embeddings' kw for backwards compatibility
   and exposes add_document(s) convenience methods.

Auto-enriches metadata at ingestion time (file_name, path, size, timestamps, page numbers, chunk_id).
"""

import os
import time
import datetime
from typing import List, Dict, Any, Optional, Type
import csv
import docx
import PyPDF2

# Embeddings
try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
except Exception:
    try:
        from langchain.embeddings import HuggingFaceEmbeddings
    except Exception:
        HuggingFaceEmbeddings = None

# Pinecone vector store
try:
    from langchain_community.vectorstores import Pinecone as LC_Pinecone
except Exception:
    LC_Pinecone = None

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from pinecone import Pinecone, ServerlessSpec

##############
#! Immortal !#
##############

DEFAULT_EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY", "pcsk_4o55Cp_Fx4jtXF6AuvZaaaqG3vhPfx4bnVCcemjcsWM67vxJad8XcBZMMpNYZF45SDEs8g")
PINECONE_ENV = os.getenv("PINECONE_ENV", "us-east-1")  # adjust if needed
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX", "immortal")


def split_text_to_docs(text: str, metadata: Dict[str, Any], chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_text(str(text))
    return [Document(page_content=chunk, metadata={**metadata}) for chunk in chunks]

def enrich_metadata(path: str, base_meta: Dict[str, Any], chunk_id: Optional[int] = None) -> Dict[str, Any]:
    try:
        stat = os.stat(path)
        size = stat.st_size
        ctime = datetime.datetime.fromtimestamp(stat.st_ctime).isoformat()
        mtime = datetime.datetime.fromtimestamp(stat.st_mtime).isoformat()
    except Exception:
        size = None
        ctime = None
        mtime = None

    enriched = dict(base_meta or {})
    enriched.update({
        "file_name": os.path.basename(path),
        "file_path": os.path.abspath(path),
        "file_extension": os.path.splitext(path)[1].lower(),
        "file_size": size,
        "created_at": ctime,
        "modified_at": mtime,
    })
    if chunk_id is not None:
        enriched["chunk_id"] = f"{os.path.basename(path)}::{chunk_id}"
    return enriched

class DocumentConverter:
    def convert(self, source_path: str, *, chunk_size=1000, chunk_overlap=100) -> List[Document]:
        raise NotImplementedError()

class CSVConverter(DocumentConverter):
    def convert(self, source_path, *, chunk_size=1000, chunk_overlap=100):
        with open(source_path, "rb") as fh:
            raw_data = fh.read()
        cleaned_data = raw_data.replace(b"\x00", b"").decode("utf-8", errors="replace")
        reader = csv.reader(cleaned_data.splitlines())
        text = "\n".join([", ".join(r) for r in reader])
        return split_text_to_docs(text, {"source": source_path, "type": "csv"}, chunk_size, chunk_overlap)

class TextConverter(DocumentConverter):
    def convert(self, source_path, *, chunk_size=1000, chunk_overlap=100):
        with open(source_path, "r", encoding="utf-8", errors="ignore") as fh:
            text = fh.read()
        return split_text_to_docs(text, {"source": source_path, "type": "text"}, chunk_size, chunk_overlap)

class PDFConverter(DocumentConverter):
    def convert(self, source_path, *, chunk_size=1000, chunk_overlap=100):
        docs: List[Document] = []
        with open(source_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for i, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if txt:
                    base_meta = {"source": source_path, "type": "pdf", "page_number": i+1}
                    docs.extend(split_text_to_docs(txt, base_meta, chunk_size, chunk_overlap))
        return docs

class DocxConverter(DocumentConverter):
    def convert(self, source_path, *, chunk_size=1000, chunk_overlap=100):
        doc = docx.Document(source_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return split_text_to_docs(text, {"source": source_path, "type": "docx"}, chunk_size, chunk_overlap)

class BaseConverter(DocumentConverter):
    def convert(self, source_path, *, chunk_size=1000, chunk_overlap=100):
        try:
            with open(source_path, "r", encoding="utf-8", errors="ignore") as fh:
                text = fh.read()
            return split_text_to_docs(text, {"source": source_path, "type": "generic"}, chunk_size, chunk_overlap)
        except Exception:
            return [Document(page_content=f"[binary file: {os.path.basename(source_path)}]", metadata={"source": source_path, "type": "binary"})]

class LoaderFactory:
    _registry: Dict[str, Type[DocumentConverter]] = {}
    @classmethod
    def register(cls, ext, converter_cls): cls._registry[ext.lower()] = converter_cls
    @classmethod
    def get_converter(cls, ext): return cls._registry.get(ext.lower(), BaseConverter)()

LoaderFactory.register(".csv", CSVConverter)
LoaderFactory.register(".txt", TextConverter)
LoaderFactory.register(".pdf", PDFConverter)
LoaderFactory.register(".docx", DocxConverter)

class ToDocument:
    @staticmethod
    def from_path(path, *, chunk_size=1000, chunk_overlap=100):
        _, ext = os.path.splitext(path)
        converter = LoaderFactory.get_converter(ext or path.split(".")[-1])
        docs = converter.convert(path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        enriched = []
        for i, d in enumerate(docs):
            enriched.append(Document(page_content=d.page_content, metadata=enrich_metadata(path, d.metadata, i)))
        return enriched

    @staticmethod
    def from_text(text, *, source="raw_text", chunk_size=1000, chunk_overlap=100):
        docs = split_text_to_docs(text, {"source": source, "type": "text"}, chunk_size, chunk_overlap)
        out = []
        for i, d in enumerate(docs):
            md = dict(d.metadata); md.setdefault("chunk_id", f"{source}::{i}")
            out.append(Document(page_content=d.page_content, metadata=md))
        return out

class PineconeStore:
    def __init__(self, embeddings=None, embedding_model=DEFAULT_EMBED_MODEL):
        if not PINECONE_API_KEY:
            raise RuntimeError("PINECONE_API_KEY not set")

        # Create Pinecone client
        self.pc = Pinecone(api_key=PINECONE_API_KEY, environment=PINECONE_ENV)

        # Set embeddings
        if embeddings is not None:
            self.embeddings = embeddings
        elif HuggingFaceEmbeddings is not None:
            self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model)
        else:
            raise RuntimeError("No embeddings available")

        # Create index if not exists
        if PINECONE_INDEX_NAME not in self.pc.list_indexes().names():
            self.pc.create_index(
                name=PINECONE_INDEX_NAME,
                dimension=len(self.embeddings.embed_query("test")),
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region=PINECONE_ENV)
            )

        # Get index reference
        self.index = self.pc.index(PINECONE_INDEX_NAME)

        # Wrap with LangChain store
        self.store = LC_Pinecone(self.index, self.embeddings.embed_query, "text")
