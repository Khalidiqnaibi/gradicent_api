"""
knowlage_base/chroma.py

Chroma-backed persistence with:
 - DocumentConverter strategies (CSV, TXT, PDF, DOCX, JSON, Generic)
 - LoaderFactory to pick converter by file extension
 - ToDocument context that coordinates conversion + chunking
 - ChromaStore accepts an 'embeddings' kw for backwards compatibility
   and exposes add_document(s) convenience methods, metadata_map persistence,
   and store (the underlying Chroma instance) for retrieval calls.

Auto-enriches metadata at ingestion time (file_name, path, size, timestamps, page numbers, chunk_id).
"""

import os
import pickle
import time
import datetime
from typing import List, Dict, Any, Optional, Type
import csv
import docx
import PyPDF2

# LangChain imports: be tolerant of slightly different versions
try:
    from langchain_community.embeddings import FastEmbedEmbeddings
except Exception:
    FastEmbedEmbeddings = None

try:
    # prefer HuggingFace wrapper if available
    from langchain_community.embeddings import HuggingFaceEmbeddings
except Exception:
    try:
        from langchain.embeddings import HuggingFaceEmbeddings
    except Exception:
        HuggingFaceEmbeddings = None

try:
    # vector store - older/newer import paths
    from langchain_community.vectorstores import Chroma
except Exception:
    try:
        from langchain.vectorstores import Chroma
    except Exception:
        Chroma = None

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

##############
#! Immortal !#
##############

DATA_FOLDER = "data_store"
CHROMA_DIR = os.path.join(DATA_FOLDER, "chroma_store")
METADATA_PATH = os.path.join(DATA_FOLDER, "metadata.pkl")

# fallback embedding model name (used if no embeddings object passed)
DEFAULT_EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

def split_text_to_docs(
    text: str,
    metadata: Dict[str, Any],
    chunk_size: int = 1000,
    chunk_overlap: int = 100
) -> List[Document]:
    """
    Split a long text into LangChain Document chunks with provided metadata.
    """
    if not isinstance(text, str):
        # protect against bytes etc.
        try:
            text = str(text, "utf-8", errors="ignore")
        except Exception:
            text = str(text)
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_text(text)
    return [Document(page_content=chunk, metadata={**metadata}) for chunk in chunks]


def enrich_metadata(path: str, base_meta: Dict[str, Any], chunk_id: Optional[int] = None) -> Dict[str, Any]:
    """Add standard metadata fields automatically for any new file."""
    try:
        stat = os.stat(path)
        size = stat.st_size
        ctime = datetime.datetime.fromtimestamp(stat.st_ctime).isoformat()
        mtime = datetime.datetime.fromtimestamp(stat.st_mtime).isoformat()
    except Exception:
        size = None
        ctime = None
        mtime = None

    enriched = dict(base_meta or {})
    enriched.update({
        "file_name": os.path.basename(path),
        "file_path": os.path.abspath(path),
        "file_extension": os.path.splitext(path)[1].lower(),
        "file_size": size,
        "created_at": ctime,
        "modified_at": mtime,
    })
    if chunk_id is not None:
        enriched["chunk_id"] = f"{os.path.basename(path)}::{chunk_id}"
    return enriched

class DocumentConverter:
    """Base class for converters. Subclasses implement convert(path) -> List[Document]."""
    def convert(self, source_path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        raise NotImplementedError()


class CSVConverter(DocumentConverter):
    def convert(self, source_path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        # Read raw bytes first to sanitize binary-like CSV files
        with open(source_path, "rb") as fh:
            raw_data = fh.read()

        # Remove NUL bytes and decode safely
        cleaned_data = raw_data.replace(b"\x00", b"").decode("utf-8", errors="replace")

        reader = csv.reader(cleaned_data.splitlines())
        lines = [", ".join(row) for row in reader]
        text = "\n".join(lines)

        base_meta = {"source": source_path, "type": "csv"}
        return split_text_to_docs(text, base_meta, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


class TextConverter(DocumentConverter):
    def convert(self, source_path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        with open(source_path, "r", encoding="utf-8", errors="ignore") as fh:
            text = fh.read()
        base_meta = {"source": source_path, "type": "text"}
        return split_text_to_docs(text, base_meta, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


class PDFConverter(DocumentConverter):
    def convert(self, source_path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        docs: List[Document] = []
        with open(source_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for i, page in enumerate(reader.pages):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt:
                    # tag pages so chunk metadata can contain page_number
                    page_text = f"--- PAGE {i+1} ---\n{txt}"
                    base_meta = {"source": source_path, "type": "pdf", "page_number": i + 1}
                    docs.extend(split_text_to_docs(page_text, base_meta, chunk_size=chunk_size, chunk_overlap=chunk_overlap))
        return docs


class DocxConverter(DocumentConverter):
    def convert(self, source_path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        doc = docx.Document(source_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paragraphs)
        base_meta = {"source": source_path, "type": "docx"}
        return split_text_to_docs(text, base_meta, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


class BaseConverter(DocumentConverter):
    """
    Fallback converter: attempts to read file as text, or if binary returns a single document with metadata.
    """
    def convert(self, source_path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        try:
            with open(source_path, "r", encoding="utf-8", errors="ignore") as fh:
                text = fh.read()
            metadata = {"source": source_path, "type": "generic_text"}
            return split_text_to_docs(text, metadata, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        except Exception:
            metadata = {"source": source_path, "type": "binary"}
            # Don't attempt to chunk binary content; store a placeholder note
            return [Document(page_content=f"[binary file stored: {os.path.basename(source_path)}]", metadata=metadata)]

class LoaderFactory:
    _registry: Dict[str, Type[DocumentConverter]] = {}

    @classmethod
    def register(cls, ext: str, converter_cls: Type[DocumentConverter]):
        cls._registry[ext.lower()] = converter_cls

    @classmethod
    def get_converter(cls, ext: str) -> DocumentConverter:
        ext = (ext or "").lower()
        if ext in cls._registry:
            return cls._registry[ext]()
        if ext.startswith(".") and ext[1:] in cls._registry:
            return cls._registry[ext[1:]]()
        return BaseConverter()

LoaderFactory.register(".csv", CSVConverter)
LoaderFactory.register("csv", CSVConverter)
LoaderFactory.register(".txt", TextConverter)
LoaderFactory.register("txt", TextConverter)
LoaderFactory.register(".pdf", PDFConverter)
LoaderFactory.register("pdf", PDFConverter)
LoaderFactory.register(".docx", DocxConverter)
LoaderFactory.register("docx", DocxConverter)


class ToDocument:
    @staticmethod
    def from_path(path: str, *, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        if not os.path.exists(path):
            raise FileNotFoundError(path)
        _, ext = os.path.splitext(path)
        converter = LoaderFactory.get_converter(ext or os.path.basename(path).split(".")[-1])
        docs = converter.convert(path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        # Ensure each doc has chunk_id metadata for traceability and enrich with file info
        enriched_docs: List[Document] = []
        for i, d in enumerate(docs):
            md = dict(d.metadata or {})
            # let converter provide page_number etc; then enrich with file-level metadata
            md = enrich_metadata(path, md, chunk_id=i)
            enriched_docs.append(Document(page_content=d.page_content, metadata=md))
        return enriched_docs

    @staticmethod
    def from_text(text: str, *, source: str = "raw_text", chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        metadata = {"source": source, "type": "text"}
        docs = split_text_to_docs(text, metadata, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        out = []
        for i, d in enumerate(docs):
            md = dict(d.metadata or {})
            md.setdefault("chunk_id", f"{source}::{i}")
            out.append(Document(page_content=d.page_content, metadata=md))
        return out


class ChromaStore:
    """
    Manages Chroma vectorstore and a persistent metadata_map file.

    Backwards compatibility:
      - Older code may call ChromaStore(embeddings=some_embeddings)
      - Newer code may call ChromaStore(persist_directory=..., embedding_model=...)
    """

    def __init__(
        self,
        embeddings: Optional[Any] = None,            # backward-compatible param
        sample_doc_path: Optional[str] = None,
        doc_loader_func: Optional[callable] = None,
        persist_directory: Optional[str] = None,
        embedding_model: str = DEFAULT_EMBED_MODEL,
    ):
        ensure_dir(DATA_FOLDER)
        self.persist_directory = persist_directory or CHROMA_DIR
        ensure_dir(self.persist_directory)

        # If the caller passed an embeddings object, use it; otherwise create one (HuggingFace preferred).
        if embeddings is not None:
            self.embeddings = embeddings
        else:
            # Try to build a huggingface embeddings instance; fall back gracefully
            if HuggingFaceEmbeddings is not None:
                try:
                    self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model)
                except Exception:
                    # last resort attempt to use FastEmbed if available
                    if FastEmbedEmbeddings is not None:
                        self.embeddings = FastEmbedEmbeddings(model_name=embedding_model)
                    else:
                        self.embeddings = None
            else:
                if FastEmbedEmbeddings is not None:
                    self.embeddings = FastEmbedEmbeddings(model_name=embedding_model)
                else:
                    self.embeddings = None

        self.sample_doc_path = sample_doc_path
        self.doc_loader_func = doc_loader_func

        # Initialize or load Chroma store
        self.store = self._init_or_load_store()
        # load metadata_map (keeps document metadata history)
        self.metadata_map = self.load_metadata_map() or {"documents": {}, "created_at": time.time()}
        # persist initial metadata if missing
        self.save_metadata_map(self.metadata_map)

    def _init_or_load_store(self):
        # If Chroma class is not available, raise early
        if Chroma is None:
            raise RuntimeError("Chroma vectorstore not found in imports. Install langchain/ langchain_community packages.")

        # If directory looks populated, open existing store; otherwise create new
        if os.path.exists(self.persist_directory) and os.listdir(self.persist_directory):
            try:
                print("Chroma store loaded from disk.")
                return Chroma(persist_directory=self.persist_directory, embedding_function=self.embeddings)
            except Exception as e:
                print("Failed to open existing Chroma store, will attempt to recreate:", e)

        # create new store: optionally seed with sample docs if provided
        if self.sample_doc_path and self.doc_loader_func:
            try:
                sample_docs = self.doc_loader_func(self.sample_doc_path)
                if not isinstance(sample_docs, list):
                    sample_docs = [sample_docs]
                chroma = Chroma.from_documents(sample_docs, embedding=self.embeddings, persist_directory=self.persist_directory)
                chroma.persist()
                # Build metadata_map
                doc_entries = {}
                for doc in sample_docs:
                    key = doc.metadata.get("source") or str(hash(doc.page_content))
                    doc_entries[str(key)] = dict(doc.metadata or {})
                self.metadata_map = {"documents": doc_entries, "created_at": time.time()}
                self.save_metadata_map(self.metadata_map)
                print("New Chroma store created with sample documents.")
                return chroma
            except Exception as e:
                print("Failed to seed Chroma from sample docs:", e)

        # empty store
        try:
            return Chroma(embedding_function=self.embeddings, persist_directory=self.persist_directory)
        except TypeError:
            # different versions have different param names
            return Chroma(persist_directory=self.persist_directory, embedding_function=self.embeddings)

    # Add single doc or multiple
    def add_document(self, doc: Document):
        """Add a single LangChain Document."""
        self.add_documents([doc])

    def add_documents(self, docs: List[Document]):
        """Add multiple Documents, persist store and update metadata map."""
        if not docs:
            print("No documents provided to add.")
            return
        # normalize to list
        normalized = []
        for d in docs:
            if isinstance(d, Document):
                # ensure page_content is str
                content = getattr(d, "page_content", "") or ""
                if isinstance(content, (bytes, bytearray)):
                    try:
                        content = content.decode("utf-8", errors="ignore")
                    except Exception:
                        content = ""
                # ensure metadata exists
                md = dict(getattr(d, "metadata", {}) or {})
                normalized.append(Document(page_content=content, metadata=md))
            else:
                raise TypeError("add_documents expects a list of langchain.schema.Document objects")

        # add to chroma
        try:
            # newer langchain may expect "documents" call or "add_documents"
            if hasattr(self.store, "add_documents"):
                self.store.add_documents(normalized)
            elif hasattr(self.store, "from_documents"):
                # fallback - not ideal but attempt
                for doc in normalized:
                    self.store.add_documents([doc])
            else:
                # final fallback - try to use constructor to merge
                raise RuntimeError("Chroma store does not expose add_documents API.")
        except Exception as e:
            print("Error adding documents to Chroma:", e)
            raise

        # update metadata map under "documents"
        doc_map = self.metadata_map.get("documents", {})
        for doc in normalized:
            key = doc.metadata.get("chunk_id") or doc.metadata.get("source") or str(hash(doc.page_content))
            doc_map[str(key)] = dict(doc.metadata or {})
        self.metadata_map["documents"] = doc_map
        self.save_metadata_map(self.metadata_map)

        try:
            if hasattr(self.store, "persist"):
                self.store.persist()
        except Exception as e:
            print("Warning: failed to persist Chroma store:", e)

        print(f"{len(normalized)} documents added to Chroma store.")

    def get_metadata_map(self) -> Dict[str, Any]:
        return self.metadata_map

    def save_metadata_map(self, metadata: Dict[str, Any]):
        ensure_dir(DATA_FOLDER)
        try:
            with open(METADATA_PATH, "wb") as f:
                pickle.dump(metadata, f)
        except Exception as e:
            print("Failed to save metadata_map:", e)

    def load_metadata_map(self) -> Optional[Dict[str, Any]]:
        if os.path.exists(METADATA_PATH):
            try:
                with open(METADATA_PATH, "rb") as f:
                    return pickle.load(f)
            except Exception as e:
                print("Failed to load metadata_map (corrupt?), starting fresh:", e)
                return {"documents": {}, "created_at": time.time()}
        # no metadata file present
        return {"documents": {}, "created_at": time.time()}

    def save_local(self):
        """Persist both store and metadata to disk."""
        try:
            if hasattr(self.store, "persist"):
                self.store.persist()
            self.save_metadata_map(self.metadata_map)
            print("Chroma store and metadata persisted to disk.")
        except Exception as e:
            print("Failed to persist Chroma store or metadata:", e)
